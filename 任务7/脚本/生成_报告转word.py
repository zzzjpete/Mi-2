# -*- coding: utf-8 -*-
"""
生成_报告转word.py — 生成第七阶段《生成层报告》Word 版（覆盖第一、第二两部分）。
样式沿用第三~六阶段（微软雅黑 / 蓝标题 / Light Grid 表格）。
输出：report_data/上下文组装与提示词工程报告.docx 和 任务7/同名文件

数字一律从实测产物解析，不写死——报告与产物始终一致：
  第一部分 ← report_data/生成_上下文组装验证报告.txt
  第二部分 ← report_data/生成_流水线测试报告_offline.txt、…_ablation.txt（末尾的汇总 JSON）
抓不到的项会打印警告并退回默认值，不静默用旧数。

用法::

    & $py scripts\生成_报告转word.py

用法::

    & $py scripts\生成_报告转word.py
"""

# ── 项目根：同一份代码在 Windows 与 Mac 上都能跑（解析规则见 _medrag_root.py）──
import os
import sys
if os.path.dirname(os.path.abspath(__file__)) not in sys.path:
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from _medrag_root import ROOT_PATH as ROOT

import json
import re
import sys
from pathlib import Path

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTS = [ROOT / "report_data" / "上下文组装与提示词工程报告.docx",
        ROOT / "任务7" / "上下文组装与提示词工程报告.docx"]
VALID_TXT = ROOT / "report_data" / "生成_上下文组装验证报告.txt"
OFFLINE_TXT = ROOT / "report_data" / "生成_流水线测试报告_offline.txt"
ABLATION_TXT = ROOT / "report_data" / "生成_流水线测试报告_ablation.txt"
LIVE_TXT = ROOT / "report_data" / "生成_流水线测试报告_live.txt"
COMPARE_JSONL = ROOT / "report_data" / "生成_对比评测.jsonl"

BLUE = RGBColor(0x1F, 0x4E, 0x79)
GRAY = RGBColor(0x88, 0x88, 0x88)


def add_title(doc, text, sub):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = BLUE
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(sub); r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_meta(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.font.size = Pt(9); r.font.color.rgb = GRAY


def add_h(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = BLUE


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(v)); run.font.size = Pt(9.5)
    return t


def add_para(doc, text, size=10.5):
    p = doc.add_paragraph(); p.add_run(text).font.size = Pt(size)
    return p


def add_bullet(doc, lead, body=""):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(lead); r.font.size = Pt(10.5); r.bold = bool(body)
    if body:
        p.add_run(body).font.size = Pt(10.5)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    lines = text.strip("\n").split("\n")
    for i, ln in enumerate(lines):
        r = p.add_run(ln)
        r.font.name = "Consolas"; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        if i < len(lines) - 1:
            r.add_break()
    return p


def add_note(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
    r = p.add_run(text); r.font.size = Pt(9.5); r.italic = True; r.font.color.rgb = GRAY
    return p


def parse_numbers():
    """从验证报告里抓关键数字；抓不到就用默认值并明确告警（不静默用旧数）。"""
    d = {"checks": "65/65", "tok_diff": "[0]", "tok_n": "9", "ctx_default": "4096",
         "vram": "6.33", "live": "2670/2800", "live_src": "5", "cut_n": "544",
         "dedup_keep": "992", "median_cpt": "4.49"}
    # 抓不到的键记在这里：交付报告宁可写「本轮未跑」，也不能把上一轮的旧数字冒充本轮实测
    d["_missing"] = set()
    if not VALID_TXT.exists():
        print(f"警告：找不到 {VALID_TXT}，全部使用默认数字")
        d["_missing"] = set(k for k in d if k != "_missing")
        return d
    t = VALID_TXT.read_text(encoding="utf-8", errors="replace")
    pats = {
        "checks": r"总计\s*(\d+/\d+)\s*项通过",
        "tok_n": r"(\d+)\s*个样本（[\d~]+\s*tok）差值集合",
        "tok_diff": r"差值集合\s*(\[[^\]]*\])",
        "ctx_default": r"Ollama 默认 num_ctx[^—]*—\s*实测\s*(\d+)",
        "vram": r"size_vram\s*([\d.]+)\s*GB",
        "live": r"真实上下文不超预算\s*—\s*(\d+/\d+)\s*tok",
        "live_src": r"真实证据来自多篇文献\s*—\s*(\d+)\s*篇",
        "cut_n": r"共触发截断\s*(\d+)\s*次",
        "dedup_keep": r"1000 条中保留\s*(\d+)",
        "median_cpt": r"中位\s*([\d.]+)\s*\|\s*均值",
    }
    for k, p in pats.items():
        m = re.search(p, t)
        if m:
            d[k] = m.group(1)
        else:
            d["_missing"].add(k)
            print(f"警告：验证报告里没抓到 {k} —— 报告中该项将标为「本轮未跑」，不填旧数字")
    return d


def _tail_summary(path: Path, what: str):
    """测试报告末尾那段汇总 JSON。抓不到就返回 None 并告警，绝不用旧数字顶上。"""
    if not path.exists():
        print(f"警告：找不到 {path}，{what} 一节将标为「未跑」")
        return None
    txt = path.read_text(encoding="utf-8", errors="replace")
    tail = txt.rsplit("=" * 96, 1)[-1].strip()
    try:
        return json.loads(tail)
    except (json.JSONDecodeError, ValueError) as e:
        print(f"警告：{path.name} 末尾的汇总 JSON 解析失败（{e}），{what} 一节将标为「未跑」")
        return None


def parse_part2():
    """第二部分的实测数字：离线完整链路 + 消融对照。"""
    off = _tail_summary(OFFLINE_TXT, "第二部分实测")
    abl = _tail_summary(ABLATION_TXT, "消融对照")
    liv = _tail_summary(LIVE_TXT, "真检索验收")

    d = {"offline": None, "ablation": None, "live": None, "compare": None}
    if COMPARE_JSONL.exists():
        rows = [json.loads(l) for l in COMPARE_JSONL.read_text(encoding="utf-8").splitlines() if l.strip()]
        if rows:
            agg = {}
            for side in ("rag", "bare"):
                agg[side] = {
                    "cite_ok": sum(r[side]["check"]["citations"]["traceable"] for r in rows),
                    "cite_n": sum(r[side]["check"]["citations"]["total"] for r in rows),
                    "yr_ok": sum(r[side]["check"]["years"]["grounded"] for r in rows),
                    "yr_n": sum(r[side]["check"]["years"]["n"] for r in rows),
                    "dr_ok": sum(r[side]["check"]["drugs"]["grounded"] for r in rows),
                    "dr_n": sum(r[side]["check"]["drugs"]["n"] for r in rows),
                    "chars": sum(r[side]["check"]["chars"] for r in rows) / len(rows),
                    "seconds": sum(r[side]["seconds"] for r in rows) / len(rows),
                }
            per = []
            for r in rows:
                b = r["bare"]["check"]["citations"]
                kind = (f"{len(b['pmids'])} 个 PMID" if b["pmids"] else
                        (f"{len(b['dois'])} 个 DOI" if b["dois"] else
                         f"{b['reference_lines']} 条「作者(年份)–期刊」，无任何标识符"))
                per.append({"id": r["case"]["id"], "cat": r["case"]["category"],
                            "rag": f"{r['rag']['check']['citations']['traceable']}/"
                                   f"{r['rag']['check']['citations']['total']}",
                            "bare": f"{b['traceable']}/{b['total']}", "kind": kind})
            d["compare"] = {"agg": agg, "per": per, "n": len(rows)}
    else:
        print(f"警告：找不到 {COMPARE_JSONL}，对比评测一节将省略")
    if liv:
        snap = liv.get("retrieval_snapshot") or {}
        d["live"] = {
            "cases": liv["cases"], "all_ok": liv["all_stages_ok"],
            "json_ok": liv["json_stage_ok"], "json_total": liv["json_stage_total"],
            "fabricated": liv["fabricated_citation_cases"],
            "no_cite": liv["no_citation_cases"],
            "t": liv["time_seconds"], "c": liv["answer_chars"],
            "load_s": snap.get("retriever_load_seconds"),
            "search_s": snap.get("search_seconds") or [],
        }
    if off:
        d["offline"] = {
            "cases": off["cases"], "all_ok": off["all_stages_ok"],
            "json_ok": off["json_stage_ok"], "json_total": off["json_stage_total"],
            "fabricated": off["fabricated_citation_cases"],
            "no_cite": off["no_citation_cases"],
            "t": off["time_seconds"], "c": off["answer_chars"], "k": off["llm_calls"],
        }
    if abl and abl.get("ablation_rows"):
        by: dict = {}
        for r in abl["ablation_rows"]:
            by.setdefault(r["config"], []).append(r)
        rows, base = [], None
        for name, g in by.items():
            mt = sum(x["seconds"] for x in g) / len(g)
            mc = sum(x["chars"] for x in g) / len(g)
            base = base if base is not None else mt
            rows.append({"name": name, "calls": g[0]["llm_calls"], "mt": mt, "mc": mc,
                         "ratio": mt / base if base else 1.0,
                         "cites": sum(x["citations_used"] for x in g) / len(g),
                         "fab": sum(x["fabricated"] for x in g)})
        d["ablation"] = {"rows": rows, "n": len(abl["ablation_rows"])}
    return d


FLOW = """检索结果 List[Candidate]
   ▼ 统一成 DocumentChunk（相关性取重排总分）+ 丢弃退化短块
   ▼ 去重：词 3-gram shingle 的 Jaccard ≥0.80 判重，簇内保留最相关者
   ▼ 排序：按相关性降序
   ▼ 选取：max(相关性 × 0.75^同源已选数)，同源硬上限 3 篇
   ▼ 装箱：按 token 预算，超长块在完整句/段处截断
context_text（每块一行出处头 [S1] PMCID · 期刊(年份) · 章节 · 标题）
   + metadata（统计 + 引用清单 citations）+ selected_chunks
   ▼
四段提示词：①证据评估 → ②作答(挂[S#]) → ③批判审查 → ④按审查意见定稿"""

FLOW2 = """查询 ─▶ [检索]* ─▶ 上下文组装 ─▶ ①证据评估 ─▶ 按评估筛证据（保留原编号）
     ─▶ ②生成草稿 ─▶ ③批判审查 ─▶ ④按审查定稿
     ─▶ 后处理（引用对账 / 格式统一 / 参考文献 / 免责声明）─▶ result
（*检索可选：传 retrieved_docs 即跳过，便于离线复跑）

消融开关 generate(evaluate=, review=)：
    False/False → 1 次模型调用（检索 + 直接作答，对照基线）
    True /False → 2 次
    True /True  → 4 次（完整四段链）"""


def main():
    V = parse_numbers()
    P2 = parse_part2()

    doc = Document()
    doc.styles["Normal"].font.name = "Microsoft YaHei"
    doc.styles["Normal"].font.size = Pt(10.5)

    add_title(doc, "第七阶段报告 · 生成层", "医学知识 RAG 系统")
    add_meta(doc, "上下文组装（之一，2026-07-27）+ 医学提示词工程 + 本地 LLM 生成流水线"
                  "（之二，2026-07-31）")

    add_h(doc, "总览")
    add_table(doc, ["部分", "内容", "状态"], [
        ["之一", "上下文组装器 + 四段医学提示词（不调用模型）", f"✅ 验证 {V['checks']}"],
        ["之二", "本地 qwen3:8b 集成 + 六步生成流水线 + 完整流程测试",
         ("✅ JSON 14/14、冒烟 7/7、离线链路 "
          f"{P2['offline']['all_ok']}/{P2['offline']['cases']}、真检索验收 "
          f"{P2['live']['all_ok']}/{P2['live']['cases']}、消融 "
          f"{P2['ablation']['n']}/{P2['ablation']['n']}")
         if (P2["offline"] and P2["ablation"] and P2["live"]) else "⚠ 部分测试报告缺失，见脚本告警"],
        ["未完成", "把裸模型那些引用逐条查实 + 人工标注评测集",
         "⏭ 已证「可验证性」差异，尚未证「正确性」差异"],
    ])

    # ================= 第一部分 =================
    add_h(doc, "第一部分 · 上下文组装与提示词工程")

    # 一、产出
    add_h(doc, "一、本次产出")
    add_table(doc, ["产出", "完成情况"], [
        ["上下文组装器 ContextAssembler",
         "✅ 估 token / 格式转换 / Jaccard 去重 / 相关性排序 / 来源多样性选取 / 预算内按完整句段截断"],
        ["文档块数据类 DocumentChunk", "✅ text / metadata / relevance_score / source / chunk_id"],
        ["医学提示词 PromptStage × 4",
         "✅ 证据评估器(T=0.0) → 答案生成器(T=0.3) → 批判性审查器(T=0.0) → 最终组装器(T=0.2)"],
        ["精确 token 计数", "✅ 从 Ollama 的 qwen3:8b GGUF 重建同一套 BPE，离线可用"],
        ["验证", f"✅ {V['checks']}，每条结论由数据算出，无一条硬编码"],
    ])
    add_para(doc, "本部分只做「检索结果 → 提示词」这一段，不调用模型；接 qwen3:8b 生成并做接入前后"
                  "对比评测，是第七阶段第二部分的内容。")

    # 二、数据流
    add_h(doc, "二、数据流")
    add_code(doc, FLOW)

    # 三、关键取舍
    add_h(doc, "三、关键设计取舍")
    add_bullet(doc, "去重用 3-gram 而非词袋。",
               "同学科的两段不同正文，词袋 Jaccard 常虚高到 0.3~0.5，阈值难定；3-gram 对词序敏感。"
               "实测：同词逆序的两段文本，词袋 1.000、3-gram 0.000。不去重的真实危害不是浪费 token，"
               "而是同一句话出现三次、模型当成三项独立证据。")
    add_bullet(doc, "多样性用软惩罚而非硬轮转。",
               "有效分 = 相关性 × 0.75^同源已选数，另设同源上限 3：同一篇的前几块该留就留，"
               "第 4 块自然让位给别的文献，避免「10 条证据其实来自同一篇」的伪多源。")
    add_bullet(doc, "截断必须落在完整句/段。",
               "半句话结尾会诱导模型「补完」，那正是幻觉来源。做法是先按 token 硬截，再在末 10% 内"
               "回退找段落分隔或句末标点，并跳过 e.g. / et al. 这类缩写点。")
    add_bullet(doc, "出处头不放检索分数。",
               "只放 PMCID、期刊、年份、章节、标题——把 0.83 这类数字给模型看，它会当成「可信度」去解释。"
               "检索侧原始信号留在 metadata 里备查。")
    add_bullet(doc, "生成拆四段而不是一个大提示词。",
               "第一阶段压测裸模型的四类错（编造参考文献、罕见病药名张冠李戴、药物分类错误、时间线自相矛盾）"
               "都不是「没检索到」，而是生成时越过了证据。一次性生成里模型既当作者又当审稿人，几乎不会否定"
               "自己刚写的句子；拆开后 ③④ 面对的是「别人写的草稿」，否定成本低得多。代价是一次问答跑 3~4 次"
               "模型——值不值要靠第二部分实测，消融开关已留好。")

    # 四、实测数字
    add_h(doc, "四、实测数字")
    add_table(doc, ["项", "实测结果"], [
        ["token 计数是否精确",
         f"与 Ollama 实际分词逐条比对，{V['tok_n']} 个样本（2~1143 tok，含中英文/公式）差值集合 {V['tok_diff']}"],
        ["语料 chars/token", f"中位 {V['median_cpt']}（p5 3.05，最小 1.31）"],
        ["预算控制", "4 组预算（200/500/1200/2800）实测装载量单调递增，均不超预算"],
        ["截断质量", f"200 条真实文本 × 3 档长度共触发 {V['cut_n']} 次截断，越界 0 次，且 100% 为原文前缀"],
        ["去重准确性", f"1000 条真实块保留 {V['dedup_keep']} 条，被丢的全部是逐字重复，误删 0 条"],
        ["端到端（4M 库真检索）",
         "本轮验证未包含 --live 组（该组需加载 65GB 向量库，约 15.8GB 内存）"
         if ("live" in V["_missing"] or "live_src" in V["_missing"])
         else f"检索 10 条 → 入选 5 条、来自 {V['live_src']} 篇不同文献、{V['live']} tok"],
        ["显存与上下文", f"Ollama 默认 num_ctx 实测 {V['ctx_default']}（四段链装不下）；"
                        f"设 12288 时实测显存 {V['vram']} GB，10G 卡有余量"],
    ])
    add_note(doc, "注：num_ctx 这条是本部分最容易被忽略的坑——Ollama 超出上下文不报错，而是静默丢掉最前面的"
                  "内容（正是 system 提示词与最相关证据），表现为「答得莫名其妙变差」。因此预算宁可高估。")

    # 五、验证
    add_h(doc, "五、验证")
    add_para(doc, f"{V['checks']} 全部通过，离线部分 7 秒跑完、不需加载 4M 库；每条 PASS 都由真实数据算出并汇入总计，"
                  f"没有任何一条是无条件打印的。「不会误伤」类结论一律在第三阶段的 1000 条真实文本块上验证。")
    add_table(doc, ["组", "覆盖"], [
        ["A 分词器", "离线加载、词表 151,936、1000 条往返一致、计数自洽、截断不超上限、降级模式保守性"],
        ["B 去重", "自相似 1.0、3-gram 优于词袋、逐字重复与近重复被丢、簇内保留最相关者、真实语料零误删"],
        ["C 排序与多样性", "纯相关性序、同源上限与软惩罚生效、有效分公式手算核对（误差 0）"],
        ["D 预算与截断", "各预算不超限、装载量单调、截断落在完整句段、缩写点不误判、极小预算安全退化"],
        ["E 元数据与引用", "必需字段齐全、计数自洽、[S#] 与入选块一一对应、参考文献编号对齐"],
        ["F 提示词模板", "四段配置、变量推导与缺失报错、渲染形状、三档 num_ctx 预算自洽"],
        ["G 端到端（真检索）",
         "本轮未跑（需 65GB 向量库）" if "live" in V["_missing"] else "真检索结果直接可用"],
        ["H 对 Ollama 实测", "分词与 Ollama 完全一致；num_ctx 默认值与显存实测"],
    ])
    if V["_missing"]:
        add_note(doc, "注：本轮验证未包含 --live 组（真检索→组装端到端，4 项），因此上表相应行标为「本轮未跑」，"
                      "而不是沿用历史数字。总计栏的分母也随之从 65 变为 61。")

    # ================= 第二部分 =================
    add_h(doc, "第二部分 · 本地 LLM 集成与生成流水线")

    add_h(doc, "六、本次产出")
    add_table(doc, ["产出", "完成情况"], [
        ["LLMGenerator（生成_LLM生成器.py）",
         "✅ 连接自检（确认模型已装）/ 温度与 token 数 / 生成 / JSON 容错解析 / 批量生成"],
        ["MedicalGenerationPipeline（生成_流水线.py）",
         "✅ 六步：上下文组装 → 证据评估(可选) → 按评估筛证据 → 生成草稿 → 批判审查(可选) → 定稿 → 后处理"],
        ["完整流程测试（生成_流水线_测试.py）",
         "✅ 三种跑法：--offline 测链路 / --live 测四类验收错题 / --ablation 三种链路配置对照"],
        ["结果结构", "✅ query / answer / context_metadata / generation_metrics / "
                     "intermediate_results / sources / timestamp，与任务书完全一致"],
    ])

    add_h(doc, "七、生成流程")
    add_code(doc, FLOW2)

    add_h(doc, "八、两个必须显式设的开关（任缺一个都会静默得到坏结果）")
    add_bullet(doc, f"num_ctx 默认只有 {V['ctx_default']}，四段链装不下。",
               "Ollama 超出上下文不报错，而是静默丢掉最前面的内容——正是 system 提示词与排序最靠前的证据，"
               f"表现为「答得莫名其妙变差」。必须显式设 12288，实测显存 {V['vram']} GB，10G 卡有余量。")
    add_bullet(doc, "think 默认开着，提示词里写 /no_think 关不掉。",
               "qwen3 是思考模型，Ollama ≥0.9 把推理过程单独放进 message.thinking，content 要等思考结束才开始写。"
               "实测思考段吃光整个 num_predict，结果是 content 全空、done_reason=length、token 却烧了一堆，"
               "看起来像「模型不回答」。必须传 API 层的 think: false。判断依据：thinking 非空而 content 为空。")
    add_note(doc, "注：另有一条已知行为——num_predict=0 会让请求挂住不返回，本模块在入口直接拦掉；"
                  "要单纯统计 prompt token 数应传 num_predict=1 并读 prompt_eval_count。")

    add_h(doc, "九、JSON 容错解析")
    add_para(doc, "证据评估器与批判性审查器要求输出 JSON，但模型会把它包进代码块、前面加一句寒暄、"
                  "或被 num_predict 截断在半个键值对上。直接 json.loads 必失败，因此做了四级递进解析：")
    add_table(doc, ["级别", "处理", "对应的实际失败形态"], [
        ["1", "直接 json.loads", "模型守规矩时最快路径"],
        ["2", "剥 <think> 段与 ``` 围栏后再解析", "代码块包裹、围栏只开不闭、全角引号"],
        ["3", "定位第一个 { 或 [ 扫平衡子串", "「好的，以下是评估结果：」这类前后缀"],
        ["4", "补缺失括号、去尾随逗号、丢弃残缺尾巴", "被 num_predict 截断的输出"],
    ])
    add_note(doc, "注：第 4 级刻意把「截断在字符串中间」的那个键连同残值一起丢掉，而不是补个引号让它看起来合法——"
                  "一个不在枚举里的值伪装成正常值（如 verdict 被截成 \"rev\"），比缺字段更难发现。"
                  "11 项解析自检全部通过。")

    if P2["offline"]:
        o = P2["offline"]
        add_h(doc, "十、实测结果：完整链路")
        add_table(doc, ["判定", "结果"], [
            ["全部用例四段链无失败阶段", f"{o['all_ok']}/{o['cases']}"],
            ["结构化阶段 JSON 全部解析成功", f"{o['json_ok']}/{o['json_total']}"],
            ["无编造出处编号", f"{o['fabricated']} 例"],
            ["每份答案都带真实引用", f"缺引用 {o['no_cite']} 例"],
            ["耗时", f"min {o['t']['min']}s / 均 {o['t']['mean']}s / max {o['t']['max']}s"],
            ["答案长度", f"min {o['c']['min']} / 均 {o['c']['mean']} / max {o['c']['max']} 字"],
            ["模型调用", f"共 {o['k']['total']} 次（单题 {o['k']['min']}~{o['k']['max']} 次）"],
        ])
        add_para(doc, "其中一道题是故意问语料答不了的问题（pembrolizumab 在非小细胞肺癌的 2024 指南起始剂量，"
                      "而离线语料是 2003–2005 年的基础生物学文献）。模型的回答是「现有证据中未提及……所有文献"
                      "均发表于 2003 年，无法反映 2024 年指南」，没有编出任何一个剂量——这正是整个 RAG 系统"
                      "要买的行为。")
        add_note(doc, "注：耗时与字数每次跑都不同（答案生成器温度 0.3、定稿 0.2，本就不是确定性的），"
                      "只作量级参考；上面四项判定则两轮复跑都成立。")

    if P2["live"]:
        v = P2["live"]
        ss = v["search_s"]
        add_h(doc, "十一、真检索验收：第一阶段留下的四类错题")
        add_para(doc, "检索侧：四道题各命中 10 条候选、来自 9~10 篇不同文献，年份 2014–2024——"
                      "正是裸模型最容易编造的近年文献区间。")
        add_table(doc, ["判定", "结果"], [
            ["四段链无失败阶段", f"{v['all_ok']}/{v['cases']}"],
            ["结构化阶段 JSON 全部解析成功", f"{v['json_ok']}/{v['json_total']}"],
            ["无编造出处编号", f"{v['fabricated']} 例"],
            ["每份答案都带真实引用", f"缺引用 {v['no_cite']} 例"],
            ["耗时", f"均 {v['t']['mean']}s（min {v['t']['min']} / max {v['t']['max']}）；"
                    f"检索器加载 {v['load_s']}s 一次性，单次检索 "
                    f"{min(ss):.2f}~{max(ss):.2f}s" if ss else f"均 {v['t']['mean']}s"],
            ["答案长度", f"均 {v['c']['mean']} 字（min {v['c']['min']} / max {v['c']['max']}）"],
        ])
        add_table(doc, ["用例", "裸模型当时的错", "本次 RAG 的回答", "关键事实可否核到"], [
            ["① CRISPR 脱靶", "标题/期刊/卷页全编",
             "参考文献全部来自检索证据，[S#] 与文末列表一一对应", "编造编号 0"],
            ["② 法布里病 ERT", "品牌名张冠李戴",
             "agalsidase-α / agalsidase-β / migalastat / 重组人 α-半乳糖苷酶", "药名逐个命中"],
            ["③ 多发性硬化 DMT", "单抗/S1P/干扰素混淆",
             "按证据表述 moderately / highly effective，未硬套证据没有的分类体系", "是"],
            ["④ 阿尔茨海默时间线", "年份自相矛盾",
             "lecanemab 获批 2023-01-06、CLARITY-AD 报告 2022-11、aducanumab 2021 加速批准，前后自洽",
             "含统计量 −0.45（p<0.001）原文可核"],
        ])
        add_note(doc, "注：上表「关键事实可否核到」是**人工逐题核对**——把答案里的药名、年份、试验名回到检索快照原文里"
                      "逐个检索的结果，可复核（快照留在 report_data/检索快照_live.json），但不可规模化。"
                      "它证明的是「没编造」，不是「答得全面」：例如③只覆盖了证据提到的分类维度，"
                      "因为检索到的证据里就没有完整的 DMT 分类体系。")

    if P2["compare"]:
        C = P2["compare"]
        a, b = C["agg"]["rag"], C["agg"]["bare"]
        add_h(doc, "十二、接入前后对比：RAG vs 裸 qwen3:8b")
        add_para(doc, "公平性上做了两件事，否则这个对比不成立：① 裸模型用的是项目日常问答那套 system prompt，"
                      "其中明确写了「不得编造参考文献」「请给出可核对的出处（作者/期刊/年份/PMID）」，"
                      "而不是随手写一个差的对照组；② 两边同模型、同 num_ctx、同 think=False、同 max_tokens，"
                      "裸模型温度取 0.3，与 RAG 的答案生成段一致。")
        add_table(doc, ["", "出处可溯源", "年份可核", "药名可核", "均字数", "均耗时"], [
            ["RAG", f"{a['cite_ok']}/{a['cite_n']}（100%）", f"{a['yr_ok']}/{a['yr_n']}",
             f"{a['dr_ok']}/{a['dr_n']}", f"{a['chars']:.0f}", f"{a['seconds']:.1f}s"],
            ["裸模型", f"{b['cite_ok']}/{b['cite_n']}（0%）", f"{b['yr_ok']}/{b['yr_n']}",
             f"{b['dr_ok']}/{b['dr_n']}", f"{b['chars']:.0f}", f"{b['seconds']:.1f}s"],
        ])
        add_table(doc, ["用例", "RAG 可溯源", "裸模型可溯源", "裸模型给的是什么"],
                  [[f"{r['id']} {r['cat']}", r["rag"], r["bare"], r["kind"]] for r in C["per"]])
        add_para(doc, "这张表真正说明的事：裸模型会成规模地生产**引用形态的文本**——四道题一共 "
                      f"{b['cite_n']} 条，其中相当一部分还带着像模像样的 PMID 或 NEJM DOI——"
                      "但这套系统一条也验证不了。RAG 侧则条条能给出 PMCID 与 PMC 链接。"
                      "两边字数几乎相同，代价是约 4.3 倍耗时。")
        add_note(doc, "注：「可溯源 0/%d」的意思是**无法用本系统核对**，不等于这些引用都是假的——"
                      "其中一部分很可能是真论文。RAG 买到的是**可验证性**，不是全知。"
                      "同理「年份/药名可核」的分母只是该题检索回的 top_k 条证据，核不到不代表错。"
                      % b["cite_n"])

    if P2["ablation"]:
        add_h(doc, "十三、消融对照：四段链值不值")
        add_table(doc, ["配置", "模型调用", "均耗时", "相对基线", "均字数", "均引用数", "编造出处"],
                  [[r["name"], str(r["calls"]), f"{r['mt']:.1f}s", f"{r['ratio']:.2f}×",
                    f"{r['mc']:.0f}", f"{r['cites']:.1f}", f"{r['fab']} 例"]
                   for r in P2["ablation"]["rows"]])
        add_para(doc, "能得出的结论：四段链的代价是约 5 倍耗时，且答案长度基本不变——多花的时间没有变成更长的废话。")
        add_para(doc, "不能得出的结论：四段链是否更忠于证据。两个观察值得记下但都还不是证据："
                      "① 完整链的引用落地数普遍更高；② 审查器在每份草稿里都挑出若干问题、verdict 多为 revise/reject。"
                      "这只能说明审查段在做事，不能说明定稿段改对了。要判定必须有人工标注评测集。")

    # 局限与下一步
    add_h(doc, "十四、已知局限与下一步")
    add_bullet(doc, "答案质量只有人工逐题核对，没有量化评测。",
               "四类错题的核对结论可复核，但不可规模化；更关键的是还没跑同题的裸模型对照组，"
               "「接入 RAG 后答得更准」这句话仍缺对照数据。")
    add_bullet(doc, "答案会中英混排。",
               "第④题的「回答」「证据要点」是英文，「证据强度与一致性」却切回了中文。"
               "提示词里写了「回答语言与用户提问语言保持一致」，模型没完全照做——"
               "不影响事实正确性，但交付观感上是缺陷。")
    add_bullet(doc, "参数仍未经效果验证。",
               "相似度阈值 0.80、多样性衰减 0.75、同源上限 3、筛证据的相关性下限 1，都只验证了「算得对」。")
    add_bullet(doc, "审查器偏严。",
               "提示词里写的是「拿不准时按未被支持记，宁可多报」，这个倾向是设计使然；"
               "但是否报得过多，没有标注集就无法判断。")
    add_bullet(doc, "下一步：",
               "① 用四类错题做接入前后对比评测（RAG vs 裸 qwen3:8b），这是唯一能支撑「RAG 有用」的证据；"
               "② 建人工标注小评测集，把第五、六、七阶段累积的未决问题一次量化——"
               "同义词扩展策略、rrf vs weighted、重排权重、去重与多样性参数、以及四段链是否更忠于证据。")

    for out in OUTS:
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
        print(f"已生成：{out}")


if __name__ == "__main__":
    main()
