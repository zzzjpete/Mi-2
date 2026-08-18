"""
生成报告.py — 生成 / 更新《RAG 数据分析与设计说明》(Word)。

阶段二收尾：把 步骤1~4 四个分析脚本的结果汇成一份交付报告。

设计：每完成一步，对应分析脚本把结果写进 report_data\stepN.json；
本脚本读取已有的 stepN.json，从头重建整份 .docx（幂等，重跑即更新）。
文风：贴近口语、简洁、有数据支撑。
"""

# ── 项目根：同一份代码在 Windows 与 Mac 上都能跑（解析规则见 _medrag_root.py）──
import os
import sys
if os.path.dirname(os.path.abspath(__file__)) not in sys.path:
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from _medrag_root import ROOT

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

REPORT_DIR = Path(os.path.join(ROOT, "report_data"))
# 只写 report_data：交付稿在 任务2\RAG数据分析与设计说明.docx，是在本脚本产出的基础上又润色过的，
# 重跑本脚本不应该把它覆盖掉。要更新交付稿请手动比对后再拷。
OUT = REPORT_DIR / "RAG数据分析与设计说明.docx"
CN_FONT = "微软雅黑"


def load(step):
    p = REPORT_DIR / f"{step}.json"
    return json.loads(p.read_text(encoding="utf-8")) if p.exists() else None


# ---------- 样式辅助 ----------
def set_cn(run, size=None, bold=None, color=None):
    run.font.name = CN_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def para(doc, text="", size=10.5, bold=False, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        set_cn(r, size=size, bold=bold, color=color)
    return p


def bullet(doc, lead, rest, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    set_cn(p.add_run("· "), size=size)
    set_cn(p.add_run(lead + "："), size=size, bold=True)
    set_cn(p.add_run(rest), size=size)
    return p


def heading(doc, text, level):
    h = doc.add_heading(level=level)
    r = h.add_run(text)
    set_cn(r, bold=True)
    return h


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
        set_cn(c.paragraphs[0].add_run(htext), size=10, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(0)
            set_cn(cells[i].paragraphs[0].add_run(str(val)), size=10)
    return t


# ---------- 封面 / 抬头 ----------
def add_header(doc):
    p = para(doc, "《RAG 数据分析与设计说明》", size=20, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    para(doc, "医学知识 RAG · 数据尽调与切块策略设计",
         size=11, color=RGBColor(0x66, 0x66, 0x66),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, f"数据源：PubMed PMC OA · oa_comm 子集（CC BY，可商用）    "
              f"更新日期：{date.today().isoformat()}",
         size=9.5, color=RGBColor(0x88, 0x88, 0x88),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "说明：本文档随开发推进逐步更新，每完成一步补充一节。",
         size=9.5, color=RGBColor(0x88, 0x88, 0x88),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)


# ---------- 第一步 ----------
def add_summary(doc, d1, d2, d3, d4):
    c = d1["completeness"]
    n = d1["n_docs"]
    heading(doc, "摘要（一页速览）", 1)
    para(doc, f"本文是医学 RAG 系统「先摸清数据、再定切块方案」的设计说明，"
              f"分析对象为本地 PubMed oa_comm 全文样本（{n:,} 篇，CC BY 可商用）。"
              f"四步走：结构分析 → 领域理解 → 长度量化 → 分割策略。")

    para(doc, "核心发现", size=11, bold=True, space_after=3)
    bullet(doc, "数据可用",
           f"字段基本零缺失、无乱码；仅摘要缺 "
           f"{c['abstract']['missing_rate_pct']}%（已定不丢弃）、PMID 缺 "
           f"{c['pmid']['missing_rate_pct']}%（用 100% 完整的 PMCID 兜底 → 溯源覆盖 100%）。")
    bullet(doc, "全文很长",
           f"token 中位 {d3['distribution']['median']:,}，"
           f"{d3['over_512_pct']}% 的文档超过 512 → 切块是硬需求。")
    bullet(doc, "结构较规整、术语密",
           f"约 {d2['full_imrad_pct']}% 有完整 IMRaD 章节；缩写密（中位每篇 "
           f"{d2['acronyms']['median_unique_per_doc']} 个）、同义与英美拼写并存 → 更靠语义检索。")
    bullet(doc, "章节本身也偏大",
           f"{d4['section_distribution']['pct_over_512']}% 的章节仍超 512 → "
           f"单纯「按章节切」不够，需在章节内再递归兜底。")

    para(doc, "设计结论（建议）", size=11, bold=True, space_after=3)
    bullet(doc, "分割策略", "章节感知 + 递归兜底两层，chunk_size≈512 / overlap≈64。")
    bullet(doc, "元数据", "每个 chunk 挂 pmid/pmcid（溯源）、journal/pub_year（过滤）、section（章节）。")
    bullet(doc, "嵌入模型", "bge-m3（多语言、长上下文）；PubMedBERT 留作后续 A/B 对照。")

    para(doc, "")
    para(doc, "一句话价值：溯源 + 过滤两项能力，正好针对准备阶段发现的"
              "「裸模型爱编造具体事实（引文/药名）」问题——这正是 RAG 要补的部分。",
         size=10.5, color=RGBColor(0x44, 0x44, 0x44))
    doc.add_page_break()


def add_step1(doc, d):
    c = d["completeness"]
    q = d["quality"]
    m = d["metadata"]
    s = d["structure_preview"]
    g = d["length_glance_chars"]
    n = d["n_docs"]

    heading(doc, "一、数据集事实（第一步：数据加载与结构分析）", 1)
    para(doc, f"这一步把本地 oa_comm 样本包整包解析成结构化表，逐字段体检，"
              f"顺带看看哪些字段将来能派上别的用场。分析对象是全文，共 {n} 篇，"
              f"全部解析成功、没有一篇失败。")

    heading(doc, "1. 数据集概览", 2)
    para(doc, f"这批是 oa_comm baseline 的第一个分包（PMC 编号最早的一段），"
              f"共 {n} 篇全文文献，来自 {m['journal_distinct']} 种期刊。"
              f"期刊清一色是开放获取刊，PLoS、BMC 系列占大头"
              f"（PLoS Biology 就有 {m['journal_top10'].get('PLoS Biology','-')} 篇），"
              f"没有 Nature、Science 这类——因为它们不在 CC BY 可商用的子集里。")
    para(doc, f"发表年份跨度 {m['year_min']}–{m['year_max']}，但高度集中在早期："
              f"光 2004 年就有 {m['year_hist'].get('2004','-')} 篇、2005 年 "
              f"{m['year_hist'].get('2005','-')} 篇，近 5 年只有 {m['year_last5_count']} 篇。"
              f"这是「第一个 ID 分包」的特点，全量语料会均衡得多，做时间过滤时要心里有数。")

    heading(doc, "2. 字段完整性", 2)
    para(doc, "整体很干净。pmcid、doi、期刊、年份、标题、正文这六项一篇都不缺；"
              "有缺口的是摘要和 PMID 两项：")
    order = [("pmcid", "PMCID"), ("pmid", "PMID"), ("doi", "DOI"),
             ("journal", "期刊 journal"), ("pub_year", "发表年 pub_date"),
             ("title", "标题 title"), ("abstract", "摘要 abstract"),
             ("body_text", "正文 body")]
    rows = [[label, f"{c[k]['present']}", f"{c[k]['missing']}",
             f"{c[k]['missing_rate_pct']}%"] for k, label in order]
    table(doc, ["字段", "有值", "缺失", "缺失率"], rows)
    para(doc, "")
    para(doc, "⚠ 摘要缺失 8.36%，超过了 1% 这条线，需要给清洗策略（见下）。"
              "另外核实过：这 253 篇的原始 XML 里压根没有 <abstract> 标签，"
              "全是 2003–2005 年早期 PLoS 文章（当年还没有结构化摘要），"
              "是真实的数据属性，不是解析漏抓。", bold=False)

    heading(doc, "3. 摘要缺失的清洗策略", 2)
    para(doc, "结论：不丢弃。原因是我们做的是「全文」RAG，内容主力是正文（100% 完整），"
              "摘要只是可选的摘要/展示字段。所以：")
    for t in [
        "· 保留这 253 篇——它们正文完整，价值不比别人低，丢了可惜；",
        "· 摘要缺失时置空并打标记，切块和检索一律走「标题 + 正文」，不依赖摘要；",
        "· 只有在需要「摘要预览」这类展示场景时，再对缺失项降级处理（例如取正文首段）。",
    ]:
        para(doc, t, space_after=2)

    heading(doc, "4. 基础质量", 2)
    para(doc, f"没有空正文，也没有编码错误（乱码）。只有两处小瑕疵："
              f"{q['short_abstract_1_100']} 篇摘要偏短（1–100 字符）、"
              f"{q['short_body_1_500']} 篇正文偏短（1–500 字符）。"
              f"数量很少，正文超短的这 {q['short_body_1_500']} 篇建议入库前人工核验或剔除。")

    heading(doc, "5. 关键字段能不能当「元数据过滤器」和「溯源链接」", 2)
    para(doc, f"能，而且基础不错。期刊和发表年都 100% 完整，"
              f"完全可以支撑「检索近 5 年某期刊的文献」这类带条件的精确检索——"
              f"机制是把 journal、pub_year 作为每个 chunk 的元数据存进向量库、检索时过滤。"
              f"要提醒的是：本样本期刊都是 OA 刊、不含 Nature，年份也偏老，"
              f"所以「近 5 年 Nature」这种具体查询要到全量、且包含目标刊时才有意义，能力本身是具备的。")
    para(doc, f"溯源方面，PMID 覆盖 {m['pmid_coverage_pct']}%，可直接拼成 PubMed 链接"
              f"（pubmed.ncbi.nlm.nih.gov/PMID）；缺的那 9% 用 PMCID 兜底"
              f"（PMCID 100% 在，可链到 PMC 全文）。两者叠加，溯源覆盖率 100%。"
              f"这一点很关键——它正是用来治「模型爱编造参考文献」的手段：每条回答都能钉回真实原文。")
    top = list(m["journal_top10"].items())
    table(doc, ["期刊（Top 10）", "篇数"], [[k, v] for k, v in top])

    heading(doc, "6. 顺带一瞥：文本长度与结构（为后续步骤铺垫）", 2)
    para(doc, f"正文明显长：字符数中位 {g['body_median']:,}、95 分位 {g['body_p95']:,}、"
              f"最长 {g['body_max']:,}，而摘要中位才 {g['abstract_median']:,} 字符。"
              f"全文这个长度，切块几乎是必然的（第三步用 token 精确量化）。"
              f"好在结构清晰：{s['docs_with_sections']}/{n} 篇有章节，"
              f"{s['docs_with_imrad_titles']} 篇带 IMRaD 式标题（方法/结果/结论等），"
              f"中位 {s['median_sections']} 个章节——这为「按章节切」提供了很好的抓手（第四步展开）。")


def add_step2(doc, d):
    tb = d["token_bands"]
    g = d["token_len_glance"]
    im = d["imrad_coverage_pct"]
    ac = d["acronyms"]

    heading(doc, "二、领域内容理解（第二步）", 1)
    para(doc, "这一步用 bge-m3 的分词器给每篇全文算了 token 长度，按短 / 中 / 长三档"
              "各抽 6 篇精读，目的是摸清这批医学文本「长什么样」——多长、结构规不规整、"
              "术语密不密——为后面的切块、提示词和评估建立心理基线。")

    heading(doc, "1. 先看长度量级（为第三步预热）", 2)
    para(doc, f"文本明显偏长：token 中位数 {g['median']:,}，95 分位 {g['p95']:,}，"
              f"最长 {g['max']:,}。关键在于——连「最短的三分之一」的门槛都要 "
              f"{tb['q33']:,} 个 token，是嵌入目标 512 的 8 倍多。也就是说几乎没有一篇文章"
              f"能整篇塞进一个 chunk，切块基本是必然的（精确分布和切割比例留到第三步）。")
    table(doc, ["档位", "token 范围", "篇数"], [
        ["短", f"≤ {tb['q33']:,}", tb["counts"]["short"]],
        ["中", f"{tb['q33']+1:,} – {tb['q66']:,}", tb["counts"]["medium"]],
        ["长", f"> {tb['q66']:,}", tb["counts"]["long"]],
    ])
    para(doc, "")
    para(doc, "抽样精读的部分样本（可见即使「短」档也多是完整文章，只有极少数是短讯）：",
         space_after=2)
    srows = []
    for b, label in [("short", "短"), ("medium", "中"), ("long", "长")]:
        for r in d["samples"][b][:2]:
            srows.append([label, r["pmcid"], r["journal"], r["year"],
                          f"{r['tokens']:,}", r["n_sections"], r["uniq_acronyms"]])
    table(doc, ["档", "PMCID", "期刊", "年", "tokens", "章节", "唯一缩写"], srows)

    heading(doc, "2. 结构：大多遵循 IMRaD", 2)
    para(doc, f"章节标题很规整。方法 {im['methods']}%、结果 {im['results']}%、"
              f"讨论 {im['discussion']}%、结论 {im['conclusions']}%、背景 {im['background']}%；"
              f"「Introduction」只有 {im['introduction']}%，是因为这些 BMC / PLoS 刊惯用"
              f"「Background」当引言。同时具备「方法 + 结果 + 讨论(或结论)」的完整 IMRaD 约 "
              f"{d['full_imrad_pct']}%。结论：约七成文章可以直接「按章节切」，"
              f"剩下三成结构不规整，需要兜底切法（第四步定）。")

    heading(doc, "3. 术语：缩写极密", 2)
    top = "、".join(f"{a}" for a, _ in ac["top20"][:12])
    para(doc, f"医学文本缩写密度很高：中位每篇 {ac['median_unique_per_doc']} 个不同缩写，"
              f"平均每千 token 出现 {ac['mean_acronyms_per_1k_tokens']} 次。高频缩写如 {top} 等"
              f"（注：其中混着 SD、CI、OR 这类统计术语，不全是医学实体，但不影响「密度高」的判断）。"
              f"影响有二：① 嵌入模型要扛得住缩写（bge-m3 没问题）；"
              f"② 用户可能用缩写、也可能用全称提问，这是检索的一个挑战点。")

    heading(doc, "4. 同一概念的多种写法（检索的隐患）", 2)
    para(doc, "同一个意思在语料里有多种表面形式，纯关键词匹配容易漏召回。举几组实测（含该词的文档数）：",
         space_after=2)
    vrows = [[v["a"], v["docs_a"], v["b"], v["docs_b"]] for v in d["variant_pairs"]]
    table(doc, ["表述 A", "文档数", "表述 B（变体）", "文档数"], vrows)
    para(doc, "")
    para(doc, "既有「正式术语 vs 通俗说法」（myocardial infarction / heart attack），"
              "也有「美式 vs 英式拼写」并存（tumor / tumour、randomized / randomised）。"
              "结论：更该依赖语义向量检索而非字面匹配——这也再次印证了选 bge-m3；"
              "后续还可加同义词 / 查询扩展进一步兜底。")

    heading(doc, "5. 内容画像（可选）", 2)
    terms = "、".join(w for w, _ in d["top_content_terms_sample400"][:12]
                     if w not in ("only", "after", "however", "during",
                                  "some", "could", "first"))
    para(doc, f"高频内容词集中在 {terms} 一带，说明这个 baseline 首包偏分子生物学 / 基因组学"
              f"（PLoS Biology、BMC Genomics / Bioinformatics 占多），纯临床内容相对少。"
              f"全量语料覆盖面会更广。这点对后续评估选题有参考：先别指望它在纯临床问题上很强。")

    heading(doc, "6. 小结", 2)
    para(doc, "一句话：文本长（必切）、结构较规整（约七成可按章节切）、缩写和同义变体多"
              "（靠语义检索 + 溯源来兜）。这三条直接决定了第三、四步的切块方案。")


def add_step3(doc, d):
    dist = d["distribution"]
    w = d["within_pct"]
    ce = d["chunk_estimate"]

    heading(doc, "三、文本特征量化分析（第三步）", 1)
    para(doc, "这一步把「全文到底有多长」用 token 量化清楚，直接回答两个问题："
              "要不要切？切多大？分词器仍用 bge-m3。")

    heading(doc, "1. Token 长度分布", 2)
    table(doc, ["统计量", "token 数"], [
        ["最小", f"{dist['min']:,}"], ["25 分位", f"{dist['p25']:,}"],
        ["中位数", f"{dist['median']:,}"], ["均值", f"{dist['mean']:,}"],
        ["75 分位", f"{dist['p75']:,}"], ["90 分位", f"{dist['p90']:,}"],
        ["95 分位", f"{dist['p95']:,}"], ["99 分位", f"{dist['p99']:,}"],
        ["最大", f"{dist['max']:,}"],
    ])
    para(doc, "")
    img = REPORT_DIR / "token_hist.png"
    if img.exists():
        doc.add_picture(str(img), width=Inches(6.3))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, f"对照任务给的判断框架——「若 95 分位约 450，则大多不用切；若 99 分位到 1200 "
              f"才有少量长文档要切」——我们这批把这个框架推到了极端：95 分位就有 "
              f"{dist['p95']:,}、99 分位 {dist['p99']:,}。也就是说，不是「少数长文档」要切，"
              f"而是「几乎全都」要切。")

    heading(doc, "2. 到底有多少需要切", 2)
    para(doc, f"按能否装进某个 chunk 大小来看（见上图右）：只有 {w['512']}% 的文档能装进 512、"
              f"{w['1024']}% 能装进 1024、{w['4096']}% 装进 4096；即便放宽到 bge-m3 的上限 8192，"
              f"也仍有 {round(100-w['8192'],1)}% 装不下。换句话说，{d['over_512_pct']}% 的文档超过 512，"
              f"「整体不切」这条路彻底排除。而且结论不是「尽量塞大」，"
              f"而是要按检索友好的小 chunk（约 512 token）来切——chunk 太大反而稀释检索精度。")

    heading(doc, "3. 切完大概多少块（给向量库规模打个底）", 2)
    rows = []
    for cfg, label in [("512/64", "512 / 64"), ("384/64", "384 / 64"),
                       ("256/32", "256 / 32")]:
        e = ce[cfg]
        rows.append([label, f"{e['total_chunks']:,}", e["mean_per_doc"],
                     e["max_per_doc"]])
    table(doc, ["chunk_size / overlap", "总 chunk 数", "平均每篇", "最多每篇"], rows)
    para(doc, "")
    para(doc, f"注意这只是 {dist['n']:,} 篇样本。全量 oa_comm 是几十万到上百万篇，"
              f"按 512/64 外推，chunk 数会到千万级——所以第四步选策略时要同时兼顾"
              f"「切得准」和「别爆量」，512/64 是个比较平衡的起点；到全量阶段，"
              f"流式建库、批量 embedding 是必须的工程前提。")

    heading(doc, "4. 小结", 2)
    para(doc, "95 / 99 分位都远超 512，切块是硬需求；推荐以约 512 token 作为 chunk 目标"
              "（检索友好）。至于「怎么切」——整体不切、滑动窗口、还是按语义章节——留到第四步"
              "结合第二步的结构结论定。")


def add_step4(doc, d):
    sd = d["section_distribution"]
    ss = d["strategy_sim"]

    heading(doc, "四、文本分割策略（第四步）", 1)
    para(doc, "综合前三步给出正式方案：数据知道能用（第一步），语言知道很密、约七成有 IMRaD "
              "结构（第二步），长度知道几乎全都超 512（第三步）。这一步先验证「按章节切够不够」，"
              "再定最终的切法、参数和每块要挂的元数据。")

    heading(doc, "1. 先验证：按章节切，够不够？", 2)
    para(doc, f"把正文按 JATS 顶层 <sec> 切开，{sd['n_sections']:,} 个章节。"
              f"章节 token：中位 {sd['median']}、均值 {sd['mean']}、90 分位 {sd['p90']:,}、"
              f"95 分位 {sd['p95']:,}。关键：只有 {sd['pct_le_512']}% 的章节能装进 512，"
              f"还有 {sd['pct_over_512']}% 仍然超过（连中位章节 {sd['median']} 都超了）。")
    img = REPORT_DIR / "section_hist.png"
    if img.exists():
        doc.add_picture(str(img), width=Inches(5.6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "结论：按章节切能对齐「背景/方法/结果/结论」的语义边界、还能给每块贴上章节标签，"
              "但章节本身也常常太大——所以单靠章节切不够，必须在章节内再兜底切一层。")

    heading(doc, "2. 推荐策略：章节感知 + 递归兜底（两层）", 2)
    for t in [
        "① 第一层·按章节切：用 JATS <sec>（或 METHODS/RESULTS 等标题）把正文分成语义段，"
        "天然对齐 IMRaD，并记录每段所属章节；",
        "② 第二层·章节内递归兜底：章节 ≤512 的整段作为一个 chunk；>512 的用 "
        "RecursiveCharacterTextSplitter（段落→句子递归），chunk_size≈512、overlap≈64 再切细；",
        "③ 无规整章节的文档（约三成）：跳过第一层，整篇直接走递归切；",
        "④ 极短文档（整篇≤512，约 0.8%）：标题+正文当一个 Document，不切，保上下文最完整。",
    ]:
        para(doc, t, space_after=2)

    heading(doc, "3. 参数与规模", 2)
    para(doc, "chunk_size 512 / overlap 64：512 落在检索友好区间、又在 bge-m3 上限内；"
              "overlap 64（约 12%）防止句子被切断、丢失跨块上下文。规模模拟如下：")
    table(doc, ["策略", "总 chunk 数", "平均每篇"], [
        ["无脑整篇定长切（对照）", f"{d['naive_wholedoc_total_chunks']:,}", "—"],
        ["章节感知 + 递归兜底（推荐）", f"{ss['total_chunks']:,}",
         ss["mean_chunks_per_doc"]],
    ])
    para(doc, "")
    para(doc, f"章节感知比无脑切多约 17%（{ss['total_chunks']:,} vs "
              f"{d['naive_wholedoc_total_chunks']:,}）——这是换取「语义边界干净 + 每块带章节标签」"
              f"付出的代价，值得。注意这只是 3028 篇样本；全量 oa_comm 数十万到上百万篇，"
              f"chunk 会到千万级，全量阶段必须流式解析建库、批量 embedding、Chroma 分批落盘。")

    heading(doc, "4. 每个 chunk 要挂的元数据（衔接过滤检索与引用溯源）", 2)
    para(doc, "切块时同步给每个 chunk 打上元数据，这是让 RAG「可过滤、可溯源」的关键：", space_after=2)
    table(doc, ["元数据字段", "用途"], [
        ["pmid / pmcid", "溯源链接：pmid 拼 PubMed，缺失时用 pmcid 拼 PMC（覆盖 100%）——治「编造引文」"],
        ["journal", "元数据过滤：按期刊筛（如某刊）"],
        ["pub_year", "元数据过滤：按年份筛（如近 5 年）"],
        ["section", "所属章节（方法/结果…），可用于加权或答案展示"],
        ["title / chunk_id", "文献标题与块定位，便于展示和调试"],
    ])

    heading(doc, "5. 补充说明", 2)
    for t in [
        "· 摘要缺失 8.36%：按第一节策略「不丢弃、置空打标、走标题+正文」，不影响全文切块；",
        "· 内容偏倚：本 baseline 首包偏分子生物 / 基因组学，纯临床内容偏少，后续评估选题需注意，"
        "可考虑扩数据到更广的 oa_comm 分包；",
        "· 同义 / 拼写变体（第二步）：靠语义向量兜底，后续可加同义词或查询扩展进一步提召回；",
        "· 嵌入模型：选 bge-m3（多语言、长上下文）；PubMedBERT 可作为医学专用对照做 A/B（留待 RAG 阶段）；",
        "· 语言：语料以英文为主，中文提问依赖 bge-m3 的跨语言检索能力。",
    ]:
        para(doc, t, space_after=2)

    heading(doc, "6. 总结 & 下一步（RAG 阶段）", 2)
    para(doc, "一句话收尾：数据干净可用（缺失已定策略）、全文必切、推荐「章节感知 + 递归兜底」"
              "（512 / 64）、每块带 pmid/pmcid/journal/pub_year/section 元数据。")
    para(doc, "据此进入 RAG 阶段的路线：按本策略切块 → bge-m3 向量化 → 灌入 Chroma → "
              "LangChain 串检索 + Qwen3-8B 生成 → 用准备阶段留下的「模型爱答错的具体事实题」"
              "做接入前后对比验收。")


def main():
    doc = Document()
    # 正文默认字体
    st = doc.styles["Normal"]
    st.font.name = CN_FONT
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)

    add_header(doc)

    d1 = load("step1")
    d2 = load("step2")
    d3 = load("step3")
    d4 = load("step4")

    if d1 and d2 and d3 and d4:
        add_summary(doc, d1, d2, d3, d4)

    if d1:
        add_step1(doc, d1)
    if d2:
        add_step2(doc, d2)
    if d3:
        add_step3(doc, d3)
    if d4:
        add_step4(doc, d4)

    doc.save(OUT)
    print("已生成：", OUT)


if __name__ == "__main__":
    main()
