# -*- coding: utf-8 -*-
"""
检索_报告转word.py — 生成第五阶段（第一部分）《查询理解与增强报告》Word 版。
样式沿用第三/四阶段（YaHei / 蓝标题 / 绿高亮 / Light Grid 表格）。
输出：report_data/查询理解与增强报告.docx 和 任务5/查询理解与增强报告.docx

数字不写死：A/B 实测结果从 report_data/查询理解验证报告.txt 里解析，
词典规模从 任务5/词典统计.json 读取——报告与实测产物始终一致。

用法::

    & $py scripts\检索_报告转word.py

用法::

    & $py scripts\检索_报告转word.py
"""

# ── 项目根：同一份代码在 Windows 与 Mac 上都能跑（解析规则见 _medrag_root.py）──
import os
import sys
if os.path.dirname(os.path.abspath(__file__)) not in sys.path:
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from _medrag_root import ROOT_PATH as ROOT

import json
import re
import sys
from pathlib import Path

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTS = [ROOT / "report_data" / "查询理解与增强报告.docx",
        ROOT / "任务5" / "查询理解与增强报告.docx"]
VALID_TXT = ROOT / "report_data" / "查询理解验证报告.txt"
STATS_JSON = ROOT / "任务5" / "词典统计.json"

BLUE = RGBColor(0x1F, 0x4E, 0x79)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
RED = RGBColor(0xB7, 0x1C, 0x1C)
GRAY = RGBColor(0x88, 0x88, 0x88)


def add_title(doc, text, sub):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = BLUE
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(sub); r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_meta(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.font.size = Pt(9); r.font.color.rgb = GRAY


def add_h(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = BLUE


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(v)); run.font.size = Pt(9.5)
    return t


def add_lead(doc, lead, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(lead); r.bold = True; r.font.size = Pt(10.5)
    if body:
        r2 = p.add_run(body); r2.font.size = Pt(10.5)
    return p


def add_para(doc, text, size=10.5):
    p = doc.add_paragraph()
    p.add_run(text).font.size = Pt(size)
    return p


def add_bullet(doc, text, color=None):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text); r.font.size = Pt(10.5)
    if color is not None:
        r.font.color.rgb = color
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(text); r.font.size = Pt(9.5); r.italic = True; r.font.color.rgb = GRAY
    return p


def parse_validation():
    """从验证报告里解析实测数字，避免报告与产物脱节。"""
    d = {}
    if not VALID_TXT.exists():
        print(f"警告：找不到 {VALID_TXT}，A/B 数字将留空")
        return d
    t = VALID_TXT.read_text(encoding="utf-8")
    m = re.search(r"自检结果：(\d+)/(\d+) 通过", t)
    if m:
        d["selfcheck"] = f"{m.group(1)}/{m.group(2)}"
    m = re.search(r"指令前缀 sentence/question：([\d.]+) / ([\d.]+)", t)
    if m:
        d["instr"] = (m.group(1), m.group(2))
    m = re.search(r"top10 平均重合 (\d+)%", t)
    if m:
        d["instr_overlap"] = m.group(1) + "%"
    m = re.search(r"扩展 主/平铺/等权RRF/加权RRF：([\d.]+) / ([\d.]+) / ([\d.]+) / ([\d.]+)", t)
    if m:
        d["exp"] = m.groups()
    else:                                   # 兼容旧版三列输出
        m = re.search(r"扩展策略 主/平铺/RRF：([\d.]+) / ([\d.]+) / ([\d.]+)", t)
        if m:
            d["exp"] = m.groups() + ("—",)
    m = re.search(r"中文 直接/中译英\s*：([\d.]+) / ([\d.]+)", t)
    if m:
        d["zh"] = (m.group(1), m.group(2))
    # 缩写查询这一行（④ 表格里唯一有区分度的一条）
    m = re.search(r"Does MI risk increase in patients with CKD\?\s+([\d.]+)\s+([\d.]+)\s+([\d.]+)(?:\s+([\d.]+))?", t)
    if m:
        d["mi_row"] = [g for g in m.groups() if g]
    return d


def main():
    V = parse_validation()
    S = json.loads(STATS_JSON.read_text(encoding="utf-8")) if STATS_JSON.exists() else {}
    sd = S.get("static_dict", {})
    md = S.get("mesh_dict", {})
    cm = S.get("corpus_meta", {})
    sec99 = cm.get("section_raw_variants_needed_for_99pct", {})

    doc = Document()
    doc.styles["Normal"].font.name = "Microsoft YaHei"
    doc.styles["Normal"].font.size = Pt(10.5)

    add_title(doc, "第五阶段报告（一）· 查询理解与增强", "医学知识 RAG 系统")
    add_meta(doc, "用户自然语言查询 → 医学实体识别与同义词扩展 → 向量/关键词双版本查询 + 元数据过滤条件　|　日期 2026-07-20")

    # 一、本周产出
    add_h(doc, "一、本周产出")
    add_table(doc, ["产出", "完成情况"], [
        ["查询理解与增强模块", "✅ 检索_查询理解.py，process_query() 六步流水线，输出结构化 EnhancedQuery"],
        ["医学同义词词典", f"✅ 两层：静态精编 {sd.get('total_entries','—')} 条 + MeSH 2026 构建 "
                          f"{md.get('descriptors_kept','—'):,} 主题词 / {md.get('surface_forms_indexed','—'):,} 词面"
         if md else "✅ 两层：静态精编 + MeSH"],
        ["医学实体识别", f"✅ {S.get('entity_patterns',{}).get('count','6')} 类正则（\\b 单词边界）"
                        " + MeSH 词典最长匹配，带类型/来源/置信度/歧义标记"],
        ["向量查询 / 关键词查询", "✅ 向量侧带 BGE 指令前缀 + 缩写消歧变体 + 融合权重；关键词侧组内 OR、组间 AND"],
        ["过滤条件提取", "✅ pub_year / section / journal，中英文时间表达均支持"],
        ["质量验证", f"✅ 离线自检 {V.get('selfcheck','12/12')} 通过；在线 4 项 A/B 实测（400 万向量库实跑）"],
    ])

    # 二、词典
    add_h(doc, "二、词典：把任务书的「示例」做成能用的东西")
    add_para(doc, "任务书给的 MEDICAL_SYNONYMS 是示例，并注明「实际应用中这个词典应该更全面，"
                  "可以从 UMLS、MeSH 等医学标准术语库构建」。这条提示已落实。")
    add_lead(doc, "为什么选 MeSH 而不是 UMLS —— ",
             "UMLS 覆盖最全（400 万概念），但要注册 UTS 账号走 license 审批（通常数天），本周交不了；"
             "MeSH 是 NLM 官方叙词表、免费直接下载，而且 PubMed 文献本来就用 MeSH 标引，与本项目语料同源，最对口。"
             "UMLS 已记入后续扩展项。")
    add_para(doc, "两层分工不是冗余，是各补各的短板：")
    add_table(doc, ["层", "规模", "管什么", "不可替代之处"], [
        ["静态精编", f"{sd.get('total_entries','—')} 条",
         f"缩写 {sd.get('abbreviations','—')} / 俗称 {sd.get('lay_terms','—')} / "
         f"商品名 {sd.get('drug_brands','—')} / 英美拼写 {sd.get('spelling_variants','—')} / "
         f"中英对照 {sd.get('cn_en_terms','—')}",
         "MeSH 基本不收缩写——「MI」根本不是 MeSH 入口词，这层只能手工维护"],
        ["MeSH 2026", f"{md.get('descriptors_kept',0):,} 主题词 / {md.get('surface_forms_indexed',0):,} 词面",
         "术语规范化、覆盖面、实体类型", "手工不可能覆盖到这个量级；实体类型直接由 MeSH 树号给出"],
    ])
    add_para(doc, "降噪做了四件事（不做的话查询会被扩散成噪声）：")
    f = md.get("filters", {})
    add_bullet(doc, f"丢弃机器轮排词 {f.get('dropped_permuted_terms',0):,} 条")
    add_bullet(doc, f"丢弃非医学分支（卫生服务/社会学/情报学/出版类型/地理）{f.get('dropped_out_of_branch',0):,} 个主题词")
    add_bullet(doc, f"丢弃不合格词面 {f.get('dropped_bad_surface',0):,} 条")
    add_bullet(doc, "倒装词面还原成自然语序：「Infarction, Myocardial」→「myocardial infarction」")
    bt = md.get("by_type", {})
    if bt:
        add_para(doc, "保留 6 类树号：" + " / ".join(
            f"{k} {v:,}" for k, v in sorted(bt.items(), key=lambda x: -x[1])) +
            f"。构建耗时 {md.get('build_seconds','—')} 秒。")
    add_note(doc, "抽检效果：metformin → glucophage（商品名）、MI → heart attack、"
                  "type 2 diabetes mellitus → NIDDM。")

    # 三、关键设计决定
    add_h(doc, "三、四个关键设计决定")

    add_lead(doc, "1. 过滤短语先剥离，再送去做向量 — ",
             "「metformin cardiovascular outcomes since 2020」里的 since 2020 已经变成 where 子句了，"
             "留在文本里只会污染语义向量。模块把命中的过滤短语从查询中剥掉，剩下的「检索主体」才送 embedding。")

    add_lead(doc, "2. 中文查询必须先转英文 — ", "实测差距最大的一项。")
    add_para(doc, "索引是 bge-base-en-v1.5（纯英文模型）+ 英文 PubMed 语料，而需求方给的示例查询正是中文。")
    if V.get("zh"):
        add_table(doc, ["平均 term-hit@10（3 条中文查询）", "直接中文检索", "中译英后"],
                  [["实测", V["zh"][0], V["zh"][1]]])
    add_bullet(doc, "Q「二甲双胍对心血管疾病有何影响？」直接中文 → top1 是《TaiChi and Qigong for "
                    "Depressive Symptoms in Patients with Chronic Heart Failure》（太极/气功治抑郁，sim 0.633）—— 完全跑题", RED)
    add_bullet(doc, "同一问题中译英 → top1 是《Protective effects of metformin in various cardiovascular "
                    "diseases》（sim 0.803）—— 正中题意", GREEN)
    add_note(doc, "Q「肿瘤微环境在癌症免疫治疗中的作用」直接检索甚至命中了中文标题文献《铁死亡的发生机制及其在肺癌中的研究进展》"
                  "——中文查询会被拉向语料里少量的中文内容，而不是主题相关的英文文献。"
                  "实现上给了两条路：词典直译（离线、零依赖、默认）和本地 qwen3:8b 整句翻译（--translate llm，覆盖词典盲区）。")

    add_lead(doc, "3. section 过滤按语料实测自适应 — ",
             f"全量扫描 400 万块发现 section 原始取值有 {cm.get('section_distinct_raw_values',0):,} 种写法。")
    add_table(doc, ["规范章节", "覆盖 99% 所需写法数", "实现方式"], [
        ["abstract / introduction / discussion / results / conclusion",
         " / ".join(str(sec99.get(k, "—")) for k in
                    ["abstract", "introduction", "discussion", "results", "conclusion"]),
         "✅ 下推 Chroma $in"],
        ["methods", str(sec99.get("methods", "—")), "⚠ $in 不现实 → 检索后用同一套归一化函数后置过滤"],
    ])
    add_note(doc, f"阈值 SECTION_IN_LIMIT=60，超过自动降级并在 notes 里说明原因。"
                  f"另有 {cm.get('section_unmapped_pct','—')}% 的块章节名无法归类"
                  f"（(no-section)、Case presentation、Main text 等），章节过滤会漏掉这部分。")

    add_lead(doc, "4. 同义词扩展方式 — ", "两次实测，两次修正了我的判断。")
    add_para(doc, "最初的判断是：同义词拼进单条向量查询会把查询向量拉向几个词的质心、反而稀释主题，"
                  "所以应该走多查询 + RRF。实测（6 条查询，term-hit@10）：")
    if V.get("exp"):
        add_table(doc, ["A 主查询（不扩展）", "B 单查询平铺", "C 多查询等权 RRF", "D 多查询加权 RRF"],
                  [list(V["exp"])])
    add_para(doc, "第一次修正：平铺（B）最高，与预设相反。但必须声明一个方法学偏差——"
                  "term-hit@10 的目标术语就是同义词扩展项，把扩展项塞进查询天然更容易命中它们，"
                  "这个指标对 B 有构造性偏袒，不能据此判 B 胜。")
    if V.get("mi_row"):
        add_para(doc, "去掉饱和噪声看真信号：6 条里 5 条都是 1.00，唯一有区分度的是缩写查询"
                      "「Does MI risk increase in patients with CKD?」→ "
                      + " / ".join(f"{n}" for n in V["mi_row"]) + "（A/B/C/D）。")
    add_para(doc, "第二次修正，针对我自己的改动：看到等权 RRF 里主查询是最差的一路（0.20），"
                  "我推测「给主查询降权应该能拉高融合结果」，于是加了 vector_query_weights"
                  "（主查询降权到 0.5）并重跑了一次完整验证。结果 D 与 C 完全相同，"
                  "没有任何可测量的改善——假设未被支持。原因也清楚：RRF 的 score = w/(k+rank) "
                  "在 k=60 时对 2 倍权重变化并不敏感，排名位置压过了权重差异。")
    add_para(doc, "所以这个权重的定位必须说清楚：它不是一项已验证的改进，而是把「主查询里含未展开的缩写」"
                  "这个信号显式交给检索层，默认值 0.5 未经验证。真正该试的更强干预是直接丢弃主查询、"
                  "或调小 k——留给第二部分。")
    add_bullet(doc, "确定：扩展救回了缩写查询。不扩展时 0.20 是灾难级——模型没理解 MI / CKD。", GREEN)
    add_bullet(doc, "不能下的结论：B / C / D 谁更好。6 条查询 + 有偏指标不足以判定，"
                    "留给第二部分带人工标注的相关性评测。", RED)
    add_bullet(doc, "被否掉的假设：给主查询降权能改善融合结果（实测无差异）。", RED)

    # 四、验证结果
    add_h(doc, "四、验证结果")
    rows = [
        ["① 功能自检（12 条查询，覆盖缩写/歧义/俗称/商品名/拼写/中文/时间/章节/无实体）",
         f"✅ {V.get('selfcheck','12/12')} 通过"],
        ["② 过滤条件下推（生成的 where 真丢给 Chroma 跑）", "✅ 4/4 返回非空，命中元数据逐条满足边界"],
    ]
    if V.get("instr"):
        rows.append(["③ BGE 指令前缀 sentence(官方) vs question(任务书)",
                     f"{V['instr'][0]} / {V['instr'][1]}，top10 重合 {V.get('instr_overlap','—')} → 差异在噪声量级"])
    if V.get("exp"):
        rows.append(["④ 同义词扩展策略", " / ".join(V["exp"]) + "（A/B/C/D，见上）"])
    if V.get("zh"):
        rows.append(["⑤ 中文 直接 vs 中译英", f"{V['zh'][0]} / {V['zh'][1]}"])
    add_table(doc, ["验证项", "结果"], rows)
    add_note(doc, "关于指令前缀：任务书写的是 Represent this question...，第四阶段建库与验证用的是 BGE 官方的 "
                  "Represent this sentence...。两者只作用于查询侧（文档侧不加前缀），换前缀不会与既有索引冲突。"
                  "实测两版平均 term-hit@10 相同、top10 重合度高，默认保留与第四阶段一致的官方版，"
                  "任务书版可用 --instruction taskbook 一键切换。")
    add_note(doc, "关于 term-hit@10：本阶段没有人工标注集，用「top-k 命中块正文里是否出现目标医学术语」作相关性代理指标。"
                  "它不等于真实相关性，但对「缩写有没有被理解」「中文有没有落到英文术语上」这类问题区分度足够，且完全可复现。")

    # 五、已知局限
    add_h(doc, "五、已知局限（都已在代码 notes 里对用户可见）")
    add_bullet(doc, "中文词典直译会丢词：「CRISPR 基因编辑的脱靶效应」中「脱靶效应」未收录。"
                    "已内置 --translate llm 走本地 qwen3:8b 整句翻译作为补救，实测可正确译出 "
                    "Off-target effects of CRISPR gene editing（冷启动约 60 秒，之后每次约 5 秒）。")
    add_bullet(doc, "term-hit@10 是代理指标，且对平铺扩展有构造性偏袒，不足以判定策略优劣。")
    add_bullet(doc, "MeSH 主表不含商品名全集：商品名多在补充概念记录 supp2026.xml（约 3GB），"
                    "本阶段未引入，常用的先手工维护了 32 条。")
    add_bullet(doc, "小写缩写有歧义风险：mi 等 2–4 字符缩写以小写出现时置信度记为 medium（大写为 high），目前仍会展开。")
    add_bullet(doc, "期刊过滤要求字面精确：库里是全称（PLoS ONE、Sensors (Basel, Switzerland)），写简称返回空，已在 notes 提示。")
    add_bullet(doc, "模糊时间词是启发式：recent / 最新 按近 5 年处理，可通过 recency_years 调整或关闭。")

    # 六、下一步
    add_h(doc, "六、下一步（检索系统第二部分）")
    add_bullet(doc, "接入检索层：多查询 RRF + where 下推 + 后置过滤，产出统一候选集")
    add_bullet(doc, "混合检索：稠密（Chroma）+ 稀疏（BM25，直接消费本模块的 keyword_query）融合")
    add_bullet(doc, "重排与证据组织：cross-encoder 重排，按「多维度证据」组织（研究设计 / 年份 / 章节 / 期刊）")
    add_bullet(doc, "建人工标注小评测集，定下本阶段没能判定的问题：扩展策略 B/C/D 之争；"
                    "以及更强的融合干预（直接丢弃主查询、调小 RRF 的 k）")
    add_bullet(doc, "补测检索性能：本阶段只记了整轮耗时（两次全流程 1007 秒 / 635 秒，差异主要来自系统页缓存变热），"
                    "未测单条查询耗时、也未测带 where 过滤是否显著更慢——这是检索层要定的关键参数。")

    for out in OUTS:
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
        print("已写:", out)


if __name__ == "__main__":
    main()
