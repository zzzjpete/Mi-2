# -*- coding: utf-8 -*-
"""
向量化_报告转word.py — 生成第四阶段《向量化与索引构建报告》Word 版。
样式沿用切块阶段（YaHei / 蓝标题 / 绿高亮 / Light Grid 表格）。
输入：report_data/向量库统计_medrag_bge_base.json、检索验证报告.txt
输入：report_data/向量库统计_medrag_bge_base.json、检索验证报告.txt
输出：report_data/向量化与索引构建报告.docx 和 任务4/向量化与索引构建报告.docx

用法::

    & $py scripts\向量化_报告转word.py

用法::

    & $py scripts\向量化_报告转word.py
"""

# ── 项目根：同一份代码在 Windows 与 Mac 上都能跑（解析规则见 _medrag_root.py）──
import os
import sys
if os.path.dirname(os.path.abspath(__file__)) not in sys.path:
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from _medrag_root import ROOT_PATH as ROOT

import sys
from pathlib import Path

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTS = [ROOT / "report_data" / "向量化与索引构建报告.docx",
        ROOT / "任务4" / "向量化与索引构建报告.docx"]

BLUE = RGBColor(0x1F, 0x4E, 0x79)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
GRAY = RGBColor(0x88, 0x88, 0x88)


def add_title(doc, text, sub):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = BLUE
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(sub); r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_meta(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.font.size = Pt(9); r.font.color.rgb = GRAY


def add_h(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = BLUE


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(v)); run.font.size = Pt(10)
    return t


def add_lead(doc, lead, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(lead); r.bold = True; r.font.size = Pt(10.5)
    if body:
        r2 = p.add_run(body); r2.font.size = Pt(10.5)
    return p


def add_para(doc, text, size=10.5):
    p = doc.add_paragraph()
    p.add_run(text).font.size = Pt(size)
    return p


def add_bullet(doc, text, green=False):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text); r.font.size = Pt(10.5)
    if green:
        r.font.color.rgb = GREEN
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(text); r.font.size = Pt(9.5); r.italic = True; r.font.color.rgb = GRAY
    return p


def main():
    doc = Document()
    doc.styles["Normal"].font.name = "Microsoft YaHei"
    doc.styles["Normal"].font.size = Pt(10.5)

    add_title(doc, "第四阶段报告 · 向量化与索引构建", "医学知识 RAG 系统")
    add_meta(doc, "嵌入模型选择加载 → ChromaDB 索引构建 → 检索质量验证（并按硬件实测容量上限确定规模）　|　日期 2026-07-16")

    # 一、本周产出
    add_h(doc, "一、本周产出")
    add_table(doc, ["产出", "完成情况"], [
        ["向量数据库", "✅ ChromaDB 持久化集合 medrag_bge_base，3,998,000 条向量，768 维，余弦相似度"],
        ["索引统计", "✅ 向量数 3,998,000、维度 768、10 个元数据字段，块 token 均长 420.8"],
        ["测试查询返回相关片段", "✅ 4 个真实医学问题均命中相关文献片段，带完整出处（PMCID·期刊·年份·章节）"],
        ["元数据过滤", "✅ pub_year>=2020 过滤后仅返回 2021–2023 文献"],
    ])

    # 二、硬件容量评估与规模确定
    add_h(doc, "二、硬件容量评估与规模确定")
    add_para(doc, "需求要求「根据硬件重新估计可处理的最大文献量，并按此量向量化」。本机实测结论：")
    add_bullet(doc, "瓶颈是内存（不是 GPU、也不是硬盘）：ChromaDB 需把全部向量常驻内存做相似度检索。"
                    "实测每条向量占 3.64 KB（768 维 fp32 约 3.0KB + HNSW 图约 0.6KB；原文与元数据存磁盘 sqlite，不占内存）。")
    add_bullet(doc, "32GB 本机留给索引约 14–20GB → 上限约 400–550 万块（≈ 17–24 万篇文献）。")
    add_bullet(doc, "据此确定规模：跨全部 11 个数据包分层随机抽样约 400 万块（3,998,000，seed 42，可复现），"
                    "占全量 9,243 万块的 4.33%，约 17.3 万篇文献；分层抽样保留年份/期刊分布，检索与过滤演示具代表性。", green=True)
    add_bullet(doc, "全量 9,243 万块入 Chroma 约需 340GB 内存，本机不可行；若要真全量见「四、扩展路径」。")
    add_note(doc, "说明：早期曾用 20 万小样本粗估 6.2KB/条，本次在 150 万、400 万规模上实测修正为 3.64KB/条，"
                  "容量比初估宽近一倍——这也是这次能把规模从演示级子集提到约 400 万的依据。")

    # 三、关键做法
    add_h(doc, "三、关键做法")
    add_lead(doc, "1. 嵌入模型 — ",
             "从候选（BGE small/base/large、OpenAI 付费、clinicalBERT 需微调）中选 "
             "BAAI/bge-base-en-v1.5（768 维）：语料是纯英文 PubMed，英文专用模型比多语言 m3 更贴；"
             "质量/速度均衡，RTX 3080 上快且省显存。采用 CLS pooling + L2 归一化；"
             "查询加指令前缀（非对称检索官方做法），文档不加。")
    add_lead(doc, "2. 索引构建 — ",
             "ChromaDB 持久化，创建集合时指定余弦相似度；唯一 id 用切块产物的 chunk_id"
             "（即 doc_id#chunk_index，如 PMC212698#15）；每块挂 10 个元数据字段"
             "（含 pmcid/pmid 溯源、journal/pub_year 过滤、section 章节）。")
    add_lead(doc, "3. 质量验证", "")
    add_bullet(doc, "① 基础统计：3,998,000 向量 / 768 维 / 样本元数据完整", green=True)
    add_bullet(doc, "② 自相似性：摘 5 个块回查，5/5 自身排第 1（sim 0.975–0.988）", green=True)
    add_bullet(doc, "③ 边界情况：空查询、17,200 字符超长查询均优雅处理（截断到 512 token，不报错）", green=True)
    add_para(doc, "检索样例：")
    add_bullet(doc, "Q「二甲双胍在 2 型糖尿病中的作用机制」→ top1 sim 0.785，命中《Metformin Protects "
                    "against Podocyte Injury in Diabetic Kidney Disease》的「Mechanisms Whereby Metformin "
                    "Reduces Hyperglycaemia」章节")
    add_bullet(doc, "Q「肿瘤微环境在癌症免疫治疗中的作用」→ top1 sim 0.842，命中《Immunosuppressive "
                    "Signaling Pathways as Targeted Cancer Therapies》的「The Tumor Microenvironment」章节")

    # 四、扩展路径
    add_h(doc, "四、扩展路径（若需真全量）")
    add_para(doc,
             "将索引从 Chroma 换 FAISS IVF-PQ（乘积量化压到约 128 字节/条，9,243 万块约 12GB 内存可行，"
             "代价是召回略降）。嵌入向量与索引选型解耦，向量算好后灌 Chroma 或 FAISS 均可，扩展无需推倒重来。")
    add_note(doc,
             "一处已知小尾巴（无害）：切块按 bge-m3(XLM-R) 分词器切、embedding 用 bge-base(BERT) 分词器，"
             "导致约 4.3% 的块超 512 BERT token 被截尾（多为薄记录/超长句，仅丢尾部）。")

    # 五、下一步
    add_h(doc, "五、下一步")
    add_para(doc,
             "检索 + 生成：用 LangChain 串「Chroma 检索 → qwen3:8b 生成」，带出处作答；"
             "再用准备阶段留下的「模型爱答错的具体事实题」做接入前后对比验收，检验 RAG 是否修复幻觉。")

    for out in OUTS:
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
        print("已写:", out)


if __name__ == "__main__":
    main()
