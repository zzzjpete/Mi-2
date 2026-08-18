# -*- coding: utf-8 -*-
"""
约束_报告转word.py — 生成第九阶段《强约束规则与幻觉抑制报告》Word 版。
样式沿用第三~八阶段（微软雅黑 / 蓝标题 / Light Grid 表格）。
输出：report_data/强约束与幻觉抑制报告.docx 和 任务9/同名文件

数字一律从实测产物解析，不写死：
  report_data/约束_对抗测试汇总.json  ← 对抗测试的全部实测数（主来源，机器读）
  report_data/约束_验证报告.txt       ← 离线验证的分组通过数
产物缺失时打印警告并在报告里标「本轮未跑」，**不静默填旧数字**。

用法::

    & $py scripts\约束_报告转word.py

用法::

    & $py scripts\约束_报告转word.py
"""

# ── 项目根：同一份代码在 Windows 与 Mac 上都能跑（解析规则见 _medrag_root.py）──
import os
import sys
if os.path.dirname(os.path.abspath(__file__)) not in sys.path:
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from _medrag_root import ROOT_PATH as ROOT

import json
import re
import sys
from pathlib import Path
from typing import Any, Dict, Optional

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTS = [ROOT / "report_data" / "强约束与幻觉抑制报告.docx",
        ROOT / "任务9" / "强约束与幻觉抑制报告.docx"]
SUMMARY_JSON = ROOT / "report_data" / "约束_对抗测试汇总.json"
VALID_TXT = ROOT / "report_data" / "约束_验证报告.txt"

BLUE = RGBColor(0x1F, 0x4E, 0x79)
GRAY = RGBColor(0x88, 0x88, 0x88)
NA = "本轮未跑"


# ---------------------------------------------------------------------------
# 排版助手（与阶段七、八同款）
# ---------------------------------------------------------------------------
def add_title(doc, text, sub):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = BLUE
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(sub); r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_meta(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.font.size = Pt(9); r.font.color.rgb = GRAY


def add_h(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = BLUE


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(v)); run.font.size = Pt(9.5)
    return t


def add_para(doc, text, size=10.5):
    p = doc.add_paragraph(); p.add_run(text).font.size = Pt(size)
    return p


def add_bullet(doc, lead, body=""):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(lead); r.font.size = Pt(10.5); r.bold = bool(body)
    if body:
        p.add_run(body).font.size = Pt(10.5)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    lines = text.strip("\n").split("\n")
    for i, ln in enumerate(lines):
        r = p.add_run(ln)
        r.font.name = "Consolas"; r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        if i < len(lines) - 1:
            r.add_break()
    return p


def add_note(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
    r = p.add_run(text); r.font.size = Pt(9.5); r.italic = True; r.font.color.rgb = GRAY
    return p


# ---------------------------------------------------------------------------
# 读实测产物
# ---------------------------------------------------------------------------
def load_summary() -> Optional[Dict[str, Any]]:
    if not SUMMARY_JSON.exists():
        print(f"警告：找不到 {SUMMARY_JSON} —— 实测各节将标为「{NA}」。"
              f"\n      先跑：约束_跑对抗测试.py")
        return None
    return json.loads(SUMMARY_JSON.read_text(encoding="utf-8"))


def load_validation() -> Dict[str, Any]:
    """从验证报告里抓「分组通过数」与「总计」。抓不到就标 NA，不编。"""
    out: Dict[str, Any] = {"total": NA, "groups": [], "seconds": NA}
    if not VALID_TXT.exists():
        print(f"警告：找不到 {VALID_TXT} —— 验证一节将标为「{NA}」")
        return out
    t = VALID_TXT.read_text(encoding="utf-8", errors="replace")
    m = re.search(r"总计\s*(\d+/\d+)\s*项通过\s*用时\s*([\d.]+)s", t)
    if m:
        out["total"], out["seconds"] = m.group(1), m.group(2)
    else:
        print("警告：验证报告里没抓到「总计 x/y 项通过」")
    for line in t.splitlines():
        g = re.match(r"^\s{2}(\d+)/(\d+)\s{3}([A-K] .+)$", line)
        if g:
            out["groups"].append((g.group(3).strip(), f"{g.group(1)}/{g.group(2)}"))
    if not out["groups"]:
        print("警告：验证报告里没抓到分组统计")
    return out


def fmt(v: Any, spec: str = "", dash: str = "—") -> str:
    if v is None:
        return dash
    if isinstance(v, bool):
        return "是" if v else "否"
    if isinstance(v, float) and spec:
        return format(v, spec)
    return str(v)


COL = {"baseline": "基线（阶段七提示词）", "prompt_only": "仅加约束提示词",
       "constrained": "约束提示词 + 校验修正"}


def main():
    s = load_summary()
    val = load_validation()
    doc = Document()
    doc.styles["Normal"].font.name = "微软雅黑"
    doc.styles["Normal"].font.size = Pt(10.5)

    meta = (s or {}).get("meta", {})
    add_title(doc, "强约束规则开发与幻觉抑制",
              "医学知识 RAG 系统 · 第九阶段报告")
    add_meta(doc, f"实测时间：{meta.get('timestamp', NA)}｜模型 {meta.get('model', NA)}｜"
                  f"对抗用例 {meta.get('n_cases', NA)} 道｜"
                  f"证据来自阶段七真实检索快照（{meta.get('snapshot_created', NA)}）")

    # ---------------------------------------------------------------- 一
    add_h(doc, "一、这一阶段解决什么问题")
    add_para(doc, "阶段七把「检索 → 上下文 → 四段提示词 → 本地 qwen3:8b」串通了，"
                  "阶段八让答案质量变得可测量。但两者都停在「建议式约束」：提示词里写着"
                  "「只能引用真实存在的编号」，却没有任何机制保证模型照做，也没有判定"
                  "「这一次到底守没守」的手段。本阶段补的正是这一层。")
    add_table(doc, ["缺口", "本阶段的做法", "代码"], [
        ["超纲问题会被顺势编一个答案", "固定拒答短语「根据现有文献无法回答此问题」写进硬约束，"
                                       "并按精确匹配统计拒答率", "约束_提示词层.py"],
        ["引用编号可能越界或写法乱", "正则提取全部编号 → 与给定范围比对 → 越界即判幻觉，"
                                     "写法不规范可确定性改回", "约束_格式校验器.py"],
        ["数据/结论可能超出证据", "答案里的数字与证据里的数字**两边同一组正则抽取后比集合**，"
                                  "查不到就报未溯源", "约束_格式校验器.py"],
        ["格式全靠模型自觉", "章节标题、缩写全称、参考文献完整性逐项判定", "约束_格式校验器.py"],
        ["发现问题也没有补救", "层 D：先确定性修正（免费），仍不合规再回灌违规清单让模型重写",
         "约束_受限流水线.py"],
        ["没有专门逼系统犯错的题", "五类对抗用例，含对照组防过度拒答", "约束_对抗测试集.py"],
    ])

    # ---------------------------------------------------------------- 二
    add_h(doc, "二、四层约束（多层次系统提示模板）")
    add_table(doc, ["层", "内容", "写在哪", "由什么兜底"], [
        ["A 硬约束", "知识库边界 / 引用来源 / 禁止编造 / 术语规范 / 输出格式（五块）",
         "系统提示词**最前面**", "层 D 的六组校验"],
        ["B 角色与写作要求", "阶段七四段各自的职责（评估员 / 作者 / 审稿人 / 定稿人）",
         "系统提示词中段", "—"],
        ["C 输出骨架", "必需章节标题作为可填模板", "用户消息", "structure.* 检查"],
        ["D 校验与修正", "生成后逐项判定，不合规触发确定性修正 → 模型重写",
         "代码（不是提示词）", "自身即判定"],
    ])
    add_note(doc, "层 A 放最前是刻意的：指令跟随对位置敏感。但要记得阶段七的坑——Ollama 超出 "
                  "num_ctx 时静默丢掉最前面的内容，所以 num_ctx=12288 与上下文预算规划必须同时成立。")
    add_para(doc, "为什么层 D 不能省：提示词是概率性的，不是编译期约束。层 A 只能提高遵守率，"
                  "层 D 才能给出「这次到底守没守」的判定——本阶段的实测正是把这两件事分别量出来。")

    # ---------------------------------------------------------------- 三
    add_h(doc, "三、对照任务书（逐条）")
    add_table(doc, ["任务书要求", "实现", "位置"], [
        ["1.a 知识库边界：答不出必须回「根据现有文献无法回答此问题」",
         "固定短语写进硬约束 1 与自检清单；无证据路径由代码直接产出该短语（阶段七那版话术里"
         "没有这句，统计时会漏算）；校验器按精确匹配判定，近义表述记 near_miss 但不算拒答",
         "约束_提示词层.py :: REFUSAL_PHRASE / no_evidence_answer"],
        ["1.b 上下文组装时给每块唯一临时编号",
         "沿用阶段七的 [S#]（组装器在 assemble_context 里分配，与参考文献一一对应）；"
         "任务书举例的 [文献#] 由校验器识别为「写法不规范但编号有效」并确定性改回",
         "生成_上下文组装.py / 约束_格式校验器.py :: _MARKER_LOOSE"],
        ["1.b 生成后用正则提取引用编号，检查是否在范围内",
         "正则提取正文全部编号（**参考文献一节不计入**，否则准确率被代码生成的列表稀释），"
         "越界即 high 违规",
         "约束_格式校验器.py :: check_citations"],
        ["1.b 无效引用或缺失引用 → 触发重试或修正",
         "层 D：① 确定性修正（删越界编号 / 改写法 / 换回文献列表）② 复检 ③ 仍不合规则把"
         "违规清单回灌给模型重写 ④ 再复检；修坏了回退",
         "约束_受限流水线.py :: enforce"],
        ["1.c 禁止编造数据、结论或细节",
         "硬约束 3 列举禁止项；可自动判定的部分由数字溯源与参考文献比对兜底",
         "约束_格式校验器.py :: check_numeric / check_references"],
        ["1.d 缩写首次出现是否给全称（术语表或正则）",
         "预定义术语表（研究方法/统计/终点/肿瘤/神经/免疫/基因编辑）+ 正则兜底未知缩写；"
         "参考文献标题里的缩写不检查",
         "约束_格式校验器.py :: TERM_GLOSSARY / check_terminology"],
        ["1.d 必需章节标题（核心答案 / 证据总结 / 参考文献）",
         "写进层 C 骨架；校验器判缺失（high）与同义改写（medium，可确定性改回）",
         "约束_格式校验器.py :: check_structure"],
        ["1.d 参考文献是否完整（至少标题、期刊、年份）",
         "逐条与权威元数据比对；区分「语料元数据本来就缺」与「模型改写」两种不完整",
         "约束_格式校验器.py :: check_references"],
        ["2 创建对抗测试用例（可用 LLM 辅助）",
         "五类共 13 道人工用例（含对照组）；另提供 --llm-assist 用本地模型起草，"
         "草稿单独存盘不自动并入（模型自己出题又自己判分没有独立性）",
         "约束_对抗测试集.py"],
        ["2 统计幻觉率、引用准确率、格式合规率",
         "同一把尺子量基线组与实验组，另给分类别与层 D 效果",
         "约束_跑对抗测试.py"],
    ])

    # ---------------------------------------------------------------- 四
    add_h(doc, "四、实测结果")
    if not s:
        add_para(doc, f"（{NA}：缺 {SUMMARY_JSON.name}）")
    else:
        sums = s["summaries"]
        order = [k for k in ("baseline", "prompt_only", "constrained") if k in sums]
        add_para(doc, f"{meta.get('n_cases')} 道对抗题 × {len(meta.get('configs', []))} 组，"
                      f"并发 {meta.get('workers')}，总耗时 {meta.get('wall_seconds')} 秒。"
                      f"第二列「仅加约束提示词」与第三列来自**同一次生成**，"
                      f"差别只在层 D 有没有介入，因此两列之间不含采样噪声。")
        rows = []
        for label, get in [
            ("幻觉率 ↓", lambda x: fmt(x["hallucination_rate"], ".4f")),
            ("引用准确率 ↑", lambda x: fmt(x["citation_accuracy"], ".4f")),
            ("格式合规率 ↑", lambda x: fmt(x["format_compliance_rate"], ".4f")),
            ("整体合规率 ↑", lambda x: fmt(x["full_compliance_rate"], ".4f")),
            # 2026-08-11 拆三态之后，标签必须跟着改：合计率里混着"完全拒答"与"部分作答"
            ("守住边界率（应拒答的题）↑", lambda x: fmt(x["refusal"]["refusal_rate"], ".4f")),
            ("　其中完全拒答 / 部分作答",
             lambda x: f"{x['refusal'].get('full_refused', '—')} / "
                       f"{x['refusal'].get('partial_answered', '—')}"),
            ("误拒率（可答题上完全拒答）↓",
             lambda x: fmt(x["refusal"]["over_refusal_rate"], ".4f")),
            ("降级率（证据充分的题）↓",
             lambda x: fmt((x.get("complete") or {}).get("downgrade_rate"), ".4f")),
            ("事实句引用覆盖 ↑", lambda x: fmt(x["citation_coverage_mean"], ".4f")),
            ("数字可溯源比例 ↑", lambda x: fmt(x["numeric_grounded_mean"], ".4f")),
            ("章节合规率 ↑", lambda x: fmt(x["structure_ok_rate"], ".4f")),
            ("术语合规率 ↑", lambda x: fmt(x["terminology_ok_rate"], ".4f")),
            ("越界编号 / 编号总数",
             lambda x: f"{x['markers_invalid']}/{x['markers_total']}"),
        ]:
            rows.append([label] + [get(sums[k]) for k in order])
        add_table(doc, ["指标"] + [COL.get(k, k) for k in order], rows)

        add_note(doc, "口径：幻觉率 = 至少命中一条 high/medium 幻觉类违规的用例数 / 用例数"
                      "（幻觉类 = 编号越界 / 参考文献编造或改写 / 数字查不到出处 / 该拒未拒）；"
                      "引用准确率为跨用例合并的 micro 口径；格式合规率只看格式类违规。"
                      "这三个数量的都是**可自动判定的部分**，不等于「答案全对」。")

        add_para(doc, "分类别（幻觉率 / 格式合规率）：")
        cat_zh = {"A_out_of_kb": "A 超出知识库", "B_induced_fabrication": "B 诱导编造数据",
                  "C_terminology": "C 术语解释", "D_fabricated_refs": "D 诱导编造参考文献",
                  "E_control": "E 对照组（本可作答）"}
        crows = []
        for cat, zh in cat_zh.items():
            row = [zh]
            for k in order:
                c = sums[k]["by_category"].get(cat)
                row.append(f"{c['hallucination_rate']:.2f} / {c['format_compliance_rate']:.2f}"
                           if c else "—")
            crows.append(row)
        add_table(doc, ["类别"] + [COL.get(k, k) for k in order], crows)

        rp = meta.get("repair")
        if rp:
            add_h(doc, "五、层 D（校验 → 修正）的效果")
            add_table(doc, ["项", "实测"], [
                ["修正前违规合计", f"{rp['violations_before']} 条"],
                ["修正后违规合计", f"{rp['violations_after']} 条"],
                ["消除", f"{rp['removed']} 条（{fmt(rp['removed_ratio'], '.2%')}）"],
                ["被修成合规的用例", "、".join(rp["cases_turned_compliant"]) or "无"],
                ["只靠确定性修正就够的用例", "、".join(rp["deterministic_only_cases"]) or "无"],
                ["动用了模型修正的用例", "、".join(rp["llm_repair_cases"]) or "无"],
            ])
            add_note(doc, "顺序是先确定性、后模型：编号写法、越界编号、章节标题、参考文献整段替换"
                          "都不需要理解内容，正则改完就对，零成本；剩下的才值得花一次模型调用。")

        add_h(doc, "六、判定")
        vr = s.get("verdicts", [])
        add_table(doc, ["判定", "结果", "依据"],
                  [[v["name"],
                    "PASS" if v["pass"] is True else ("N/A" if v["pass"] is None else "FAIL"),
                    v["detail"]] for v in vr])
        n_pass = sum(1 for v in vr if v["pass"] is True)
        n_apply = sum(1 for v in vr if v["pass"] is not None)
        add_para(doc, f"合计 {n_pass}/{n_apply} 项通过"
                      + (f"（另有 {len(vr) - n_apply} 项本轮没有对应数据，判不了；"
                         f"记 N/A 而不是 FAIL）" if n_apply != len(vr) else ""))

    # ---------------------------------------------------------------- 七
    add_h(doc, "七、离线验证（不调用模型）")
    add_para(doc, f"约束_验证.py：总计 {val['total']} 项通过，用时 {val['seconds']} 秒。"
                  f"每条 PASS/FAIL 都由实际算出的变量比对得出，没有一条是无条件打印。")
    if val["groups"]:
        add_table(doc, ["分组", "通过"], [[g, n] for g, n in val["groups"]])

    # ---------------------------------------------------------------- 八
    add_h(doc, "八、已知局限（如实说明）")
    add_bullet(doc, "只能判定可自动判定的部分。", "措辞谨慎的事实性错误（证据说「可能相关」、"
               "答案写「可能有效」但方向反了）、逻辑错误、临床上的不当推荐，本模块测不到。"
               "「幻觉率 0」只意味着这四类可判定信号没有命中。")
    add_bullet(doc, "数字溯源有已知假阳性。", "模型把 12.63% 四舍五入成 12.6%、做了单位换算，"
               "都会被记成「证据里查不到」。所以这一项按 medium 报并列出原文供人核。")
    add_bullet(doc, "事实句覆盖率是启发式。", "程序分不出「断言句」和「转折句」，阈值 0.8 留得宽。")
    add_bullet(doc, "术语表求精不求全。", "表里没有的缩写只走正则兜底、只报 low；"
               "CRISPR/Cas9/DNA 这类当专名用的词已排除，否则每份答案都会被判「缺全称」。")
    add_bullet(doc, "用例规模小。", "13 道题、单次采样，②答案生成段温度 0.3，"
               "重跑数字会有波动；结论应看方向与量级，不要引用小数点后两位。")
    add_bullet(doc, "拒答率高不等于系统更好。", "所以必须同时看对照组的误拒率——"
               "一个「什么都拒答」的系统在 A~D 类上能拿满分。")

    for out in OUTS:
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
        print(f"已生成：{out}")


if __name__ == "__main__":
    main()
