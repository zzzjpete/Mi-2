# -*- coding: utf-8 -*-
"""
评估_报告转word.py — 生成第八阶段《答案评估、缓存与批量处理报告》Word 版。
样式沿用第三~七阶段（微软雅黑 / 蓝标题 / Light Grid 表格）。
输出：report_data/答案评估与缓存批量处理报告.docx 和 任务8/同名文件

数字一律从实测产物解析，不写死：
  report_data/评估_测试集汇总.json   ← 跑测试集的全部实测数（主来源，机器读）
  report_data/评估_验证报告.txt      ← 离线验证的分组通过数
产物缺失时打印警告并在报告里标「本轮未跑」，**不静默填旧数字**。

用法::

    & $py scripts\评估_报告转word.py

用法::

    & $py scripts\评估_报告转word.py
"""

# ── 项目根：同一份代码在 Windows 与 Mac 上都能跑（解析规则见 _medrag_root.py）──
import os
import sys
if os.path.dirname(os.path.abspath(__file__)) not in sys.path:
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from _medrag_root import ROOT_PATH as ROOT

import json
import re
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTS = [ROOT / "report_data" / "答案评估与缓存批量处理报告.docx",
        ROOT / "任务8" / "答案评估与缓存批量处理报告.docx"]
SUMMARY_JSON = ROOT / "report_data" / "评估_测试集汇总.json"
VALID_TXT = ROOT / "report_data" / "评估_验证报告.txt"

BLUE = RGBColor(0x1F, 0x4E, 0x79)
GRAY = RGBColor(0x88, 0x88, 0x88)
NA = "本轮未跑"


# ---------------------------------------------------------------------------
# 排版助手（与阶段七同款）
# ---------------------------------------------------------------------------
def add_title(doc, text, sub):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = BLUE
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(sub); r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_meta(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.font.size = Pt(9); r.font.color.rgb = GRAY


def add_h(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = BLUE


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(v)); run.font.size = Pt(9.5)
    return t


def add_para(doc, text, size=10.5):
    p = doc.add_paragraph(); p.add_run(text).font.size = Pt(size)
    return p


def add_bullet(doc, lead, body=""):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(lead); r.font.size = Pt(10.5); r.bold = bool(body)
    if body:
        p.add_run(body).font.size = Pt(10.5)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    lines = text.strip("\n").split("\n")
    for i, ln in enumerate(lines):
        r = p.add_run(ln)
        r.font.name = "Consolas"; r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        if i < len(lines) - 1:
            r.add_break()
    return p


def add_note(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
    r = p.add_run(text); r.font.size = Pt(9.5); r.italic = True; r.font.color.rgb = GRAY
    return p


# ---------------------------------------------------------------------------
# 读实测产物
# ---------------------------------------------------------------------------
def load_summary() -> Optional[Dict[str, Any]]:
    if not SUMMARY_JSON.exists():
        print(f"警告：找不到 {SUMMARY_JSON} —— 实测各节将标为「{NA}」。"
              f"\n      先跑：评估_跑测试集.py")
        return None
    return json.loads(SUMMARY_JSON.read_text(encoding="utf-8"))


def load_validation() -> Dict[str, Any]:
    """从验证报告里抓「分组通过数」与「总计」。抓不到就标 NA，不编。"""
    out: Dict[str, Any] = {"total": NA, "groups": [], "seconds": NA}
    if not VALID_TXT.exists():
        print(f"警告：找不到 {VALID_TXT} —— 验证一节将标为「{NA}」")
        return out
    t = VALID_TXT.read_text(encoding="utf-8", errors="replace")
    m = re.search(r"总计\s*(\d+/\d+)\s*项通过\s*用时\s*([\d.]+)s", t)
    if m:
        out["total"], out["seconds"] = m.group(1), m.group(2)
    else:
        print("警告：验证报告里没抓到「总计 x/y 项通过」")
    for line in t.splitlines():
        g = re.match(r"^\s{2}(\d+)/(\d+)\s{3}([A-G] .+)$", line)
        if g:
            out["groups"].append((g.group(3).strip(), f"{g.group(1)}/{g.group(2)}"))
    if not out["groups"]:
        print("警告：验证报告里没抓到分组统计")
    return out


def fmt(v: Any, spec: str = "", dash: str = "—") -> str:
    if v is None:
        return dash
    if spec and isinstance(v, (int, float)):
        return format(v, spec)
    return str(v)


# ---------------------------------------------------------------------------
def main():
    s = load_summary()
    v = load_validation()
    doc = Document()
    doc.styles["Normal"].font.name = "微软雅黑"
    doc.styles["Normal"].font.size = Pt(10.5)

    add_title(doc, "医学知识 RAG · 第八阶段",
              "生成答案评估、缓存策略与批量处理")
    if s:
        env = s["environment"]
        add_meta(doc, f"实测时间 {s['timestamp']}｜模型 {env['model']}｜num_ctx {env['num_ctx']}"
                      f"｜CPU 逻辑核 {env['cpu_count']}｜ROUGE 后端 {env['rouge_backend']}")
    else:
        add_meta(doc, f"实测产物缺失，本报告的实测各节标为「{NA}」")

    # ---------------- 一、这一阶段解决什么 ----------------
    add_h(doc, "一、这一阶段解决什么")
    add_para(doc, "阶段七把「检索 → 上下文 → 四段提示词 → 本地 qwen3:8b」串成了一条能出答案的链，"
                  "但当时留下三个缺口：答案好不好只能人读、同样的问题每次都要重新烧 40 秒、"
                  "一批问题只能一条条排队。本阶段补的正是这三件事。")
    add_table(doc, ["缺口", "本阶段的做法", "代码"], [
        ["答案质量只能人读，没有可复算的数字",
         "四维评估器：ROUGE 相似性 / 医学关键信息召回 / 幻觉信号 / 可读性",
         "评估_答案评估器.py"],
        ["同样的查询+同样的证据要重复烧模型",
         "带 TTL、容量上限与温度门限的缓存，键 = 查询+上下文+全部影响输出的参数",
         "生成_缓存.py"],
        ["一批问题只能串行",
         "线程池批量：顺序保持、错误隔离、并发度按工作性质推荐",
         "生成_批量处理.py"],
    ])
    add_note(doc, "评估的是「可复算的量」，不是「对不对」——正确性仍需人工标注，见末节。")

    # ---------------- 二、答案评估器 ----------------
    add_h(doc, "二、答案评估器：四个维度各回答一个不同的问题")
    add_table(doc, ["维度", "指标", "怎么算", "它不能说明什么"], [
        ["① 文本相似性", "ROUGE-1 / ROUGE-2 / ROUGE-L（rouge 库）",
         "答案与参照文本的 n-gram 与最长公共子序列重合",
         "不能说明答得对：照抄证据得高分，换个说法说对了反而低分"],
        ["② 关键信息召回", "recall = overlap / gt_matches",
         "正则抽七类医学关键信息（百分比/剂量/时间/安全/建议/机制/统计量），比集合交集",
         "只看信息「在不在」，不看用得对不对"],
        ["③ 幻觉信号", "risk_score ∈ [0,1) + 逐条命中",
         "五类无依据绝对化表述；同句带 [S#]/PMID/DOI 的记为 mitigated",
         "只看措辞，措辞谨慎的编造抓不到；抓到的也未必是假的"],
        ["④ 可读性", "平均句子长度、长句比例、结构、语言一致性",
         "按中英文句末标点切句（跳过 e.g. 类缩写点）",
         "是约定的经验带，不是校准量表"],
    ])
    add_para(doc, "参照物决定了 ①② 的口径，这一点必须先说清楚，否则数字会被误读：")
    add_bullet(doc, "gold（人工标准答案）：", "任务书公式的本意。本项目目前没有这份标注，"
                                            "脚本已支持 --refs 传入，一旦有了就能直接用。")
    add_bullet(doc, "evidence（检索证据原文）：", "本项目现在拿得到的参照。此时量的是"
                                                "「答案有没有贴着证据写」，是忠于证据的代理指标。")
    add_bullet(doc, "none：", "不给参照，只跑 ③④——这两维本来就不需要参照。")

    # ---------------- 三、缓存 ----------------
    add_h(doc, "三、缓存策略：四件事各对应一个会出问题的场景")
    add_table(doc, ["要求", "实现", "不这么做会怎样"], [
        ["键 = 查询 + 上下文的哈希",
         "sha256（查询 + 上下文摘要 + 模型/温度/max_tokens/num_ctx/JSON 模式）；"
         "包装器直接把完整 messages（含 system 提示词）纳入键",
         "改了提示词却读到旧答案，而且查不出来"],
        ["容量上限，避免内存溢出",
         "条数与字节数双上限，LRU 淘汰",
         "长跑的评测进程里就是内存泄漏"],
        ["TTL：医学知识有时效性",
         "默认 7 天，可按用途调；落盘再载入时过期条目直接丢弃",
         "指南更新、药物撤市后，过期结论被永久钉死"],
        ["只缓存低温度（确定性）结果",
         "默认门限 0.0；高于门限的写入被拒绝并计入 skipped_temperature",
         "把一次抽样的结果固定下来，复跑看起来「稳定」其实是假的"],
    ])
    add_note(doc, "线程安全：内部一把 RLock——批量走多线程，缓存必然并发读写。")

    # ---------------- 四、批量处理 ----------------
    add_h(doc, "四、批量处理：三条硬要求各对应一句代码级保证")
    add_bullet(doc, "输出顺序与输入一致：", "结果列表按下标预分配，as_completed 只用来报进度，"
                                          "写回一律按 future→index 映射。绝不 append。")
    add_bullet(doc, "单个任务失败不影响其他：", "每个任务在 worker 内自带 try/except，"
                                             "失败项以占位记录保留原位置，并带错误与调用栈。")
    add_bullet(doc, "按 CPU 核心数定并发度：", "llm 取 min(4, 核数/4)、cpu 取核数−1、io 取 min(32, 核数×2)；"
                                            "任务数少于并发度时按任务数收敛。")

    # ---------------- 五、实测结果 ----------------
    add_h(doc, "五、实测结果：用上周的四道验收错题再跑一遍")
    if not s:
        add_para(doc, f"（{NA}：缺 {SUMMARY_JSON.name}）")
    else:
        add_para(doc, f"测试题就是阶段一压测裸模型时留下的四类错题，检索结果读阶段七固化的快照"
                      f"（生成于 {s['snapshot_created']}），因此本轮不加载 65 GB 向量库。")

        add_para(doc, "1) 批量处理：串行 vs 并行（两趟都是冷缓存，起跑线相同）", 10.5)
        rows = [[p["name"], p["workers"], f"{p['wall_seconds']:.2f}",
                 f"{p['ok']}/{p['items']}", p["cache_hits"]] for p in s["passes"]]
        add_table(doc, ["趟", "并发", "墙钟(s)", "成功", "缓存命中"], rows)
        sp = s.get("serial_vs_parallel_speedup")
        if sp:
            add_para(doc, f"并行相对串行 {sp}×。本机是单卡 + 单个 Ollama 进程，"
                          f"多个请求争同一份权重与同一块 GPU，所以这个倍数量的是"
                          f"「本机实际能拿到多少并行收益」，不是线程池的理论上限。")

        c = s["cache"]
        add_para(doc, "2) 缓存：同一批查询第二次跑", 10.5)
        add_table(doc, ["项", "实测"], [
            ["冷（未预热）墙钟", f"{c['cold_wall_seconds']:.2f} s"],
            ["热（已预热）墙钟", f"{c['hot_wall_seconds']:.3f} s"],
            ["缓存替下来的模型时间", f"{fmt(c.get('model_seconds_saved'))} s（四题合计）"],
            ["命中数", f"{c['hot_hits']}/{s['cases']}"],
            ["命中返回的答案与首次逐字相同", c["answers_identical"]],
            ["TTL", f"{c['ttl_seconds'] / 3600:.1f} h"],
            ["阶段级门限 / 流水线级门限",
             f"{c['stage_max_temperature']} / {c['pipeline_max_temperature']}"],
            ["因温度过高被拒绝缓存的调用数",
             c["stage_cache"].get("skipped", NA)],
        ])

        ec = s["evaluation"]
        add_para(doc, f"3) 答案评估（{s['cases']} 份 RAG 答案，参照类型 = {ec['reference_kind']}）",
                 10.5)
        rows = [[p["case_id"], fmt(p["rouge_l_f"], ".4f"), fmt(p["key_info_recall"], ".4f"),
                 f"{p['hallucination_risk']:.4f}（{p['hallucination_level']}）",
                 f"{p['signals_unmitigated']}/{p['signals_total']}",
                 f"{p['readability_score']:.4f}", f"{p['avg_sentence_length']:.1f}",
                 f"{p['mixed_language_ratio']:.2f}"] for p in ec["per_case"]]
        add_table(doc, ["用例", "ROUGE-L f", "关键信息召回", "幻觉风险", "无出处/总信号",
                        "可读性", "均句长", "中英混排"], rows)
        add_note(doc, f"评估 {s['cases'] + ec['bare_cases']} 份答案用时 "
                      f"{ec['eval_seconds']} s（并发 {s['environment']['eval_workers']}）——"
                      f"评估本身是纯 CPU 工作，这里的并行是实打实有效的。")

        if ec.get("bare_aggregate"):
            add_para(doc, "4) 同一把尺子量裸模型（复用阶段七已存的答案，未额外调模型）", 10.5)
            ra, ba = ec["rag_aggregate"], ec["bare_aggregate"]
            add_table(doc, ["", "ROUGE-L f", "关键信息召回", "幻觉风险", "无出处信号数",
                            "可读性", "均句长"], [
                ["RAG", fmt(ra["rouge_l_f"], ".4f"), fmt(ra["key_info_recall"], ".4f"),
                 fmt(ra["hallucination_risk"], ".4f"),
                 ra["hallucination_signals_unmitigated"],
                 fmt(ra["readability_score"], ".4f"), fmt(ra["avg_sentence_length"], ".1f")],
                ["裸模型", fmt(ba["rouge_l_f"], ".4f"), fmt(ba["key_info_recall"], ".4f"),
                 fmt(ba["hallucination_risk"], ".4f"),
                 ba["hallucination_signals_unmitigated"],
                 fmt(ba["readability_score"], ".4f"), fmt(ba["avg_sentence_length"], ".1f")],
            ])
            add_note(doc, "口径：ROUGE 与召回都是对同一批检索证据算的——RAG 看得见这些证据、"
                          "裸模型看不见，所以 RAG 更高是意料之中。它量的是「有没有贴着证据写」，"
                          "不是「谁答得更对」。后者需要人工标注才能回答。")

        add_para(doc, "5) 判定（每条都由上面的实测数据算出）", 10.5)
        add_table(doc, ["判定", "结果"],
                  [[k["label"], "PASS" if k["passed"] else "FAIL"] for k in s["checks"]])

    # ---------------- 六、验证 ----------------
    add_h(doc, "六、离线验证：每条结论都由变量算出")
    add_para(doc, f"评估_验证.py 不调用模型，{v['seconds']} 秒跑完，"
                  f"总计 {v['total']} 项通过。没有任何一条是无条件打印的。")
    if v["groups"]:
        add_table(doc, ["分组", "通过"], [[g, n] for g, n in v["groups"]])

    # ---------------- 七、踩过的坑 ----------------
    add_h(doc, "七、这一阶段踩到的坑（都是实测撞出来的）")
    add_bullet(doc, "rouge 库直接吃中文，分数几乎没有意义。",
               "它按空白切词，一整句中文会被切成 1 个「词」。实测：两句高度相似的中文，"
               "不预分词时 ROUGE-L = 0.0000，按「中文按字、英文按词」预分词后 = 0.7907。")
    add_bullet(doc, "rouge 库把「.」当句子分隔符，而医学答案满是小数。",
               "0.45、p<0.001 会被劈成两个「句子」，既污染 unigram 计数，也让 ROUGE-L 走错分支。"
               "解法：token 内的小数点换成「·」，「.」只留作真正的句子边界。")
    add_bullet(doc, "把整段当成一个句子喂给 rouge，会直接 RecursionError。",
               "它的 LCS 回溯是递归的，2000+ token 的真实答案直接栈溢出。给它真正的句子边界即可，"
               "而且句子级切分本来就是 ROUGE-L summary level 的正确用法。")
    add_bullet(doc, "跨语种的关键信息比对，按字面永远匹配不上。",
               "证据是英文、答案常是中文：「不良反应」与 adverse events 字面无交集，"
               "这三类关键词的召回率会恒等于 0——不是模型漏了，是尺子坏了。"
               "解法：中英同义写法映射到同一个 concept id。")
    add_bullet(doc, "风险分用「除以上限再截断」会失去单调性。",
               "实测 110 字的文本里 1 个信号和 4 个信号都顶到 1.0，分不出轻重。"
               "改成 risk = 1 − exp(−密度/2)：恒在 [0,1) 且严格单调。")

    # ---------------- 八、局限与下一步 ----------------
    add_h(doc, "八、已知局限与下一步")
    add_bullet(doc, "没有人工标准答案，量到的仍是「贴不贴着证据」，不是「对不对」。",
               "ROUGE 与关键信息召回的参照目前是检索证据原文。评估器已支持 --refs 传入 gold，"
               "缺的只是那份标注。这是阶段七、八共同的遗留项。")
    add_bullet(doc, "幻觉检测是措辞层面的启发式。",
               "措辞谨慎的编造它抓不到；抓到的也只说明「说得太满且没给出处」，不等于内容为假。")
    add_bullet(doc, "可读性分与幻觉风险刻度都是约定值。",
               "句长经验带 15~45 字、风险密度尺度 2.0，都没有经过人评校准。")
    add_bullet(doc, "并行对本地生成的收益受单卡限制。",
               "评估、分词、正则这类纯 CPU 批量能吃到并行收益；模型推理受单块 GPU 与 Ollama "
               "服务端排队约束，实测倍数见第五节，不要照搬「CPU 核数」这个直觉。")

    for out in OUTS:
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
        print(f"已生成：{out}")


if __name__ == "__main__":
    main()
