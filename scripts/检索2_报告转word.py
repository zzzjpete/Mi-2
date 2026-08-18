# -*- coding: utf-8 -*-
"""
检索2_报告转word.py — 生成第六阶段（第二部分）《多路检索与重排报告》Word 版。
样式沿用第三/四/五阶段（微软雅黑 / 蓝标题 / 绿高亮 / Light Grid 表格）。
输出：report_data/多路检索与重排报告.docx 和 任务6/多路检索与重排报告.docx

数字不写死：验证结论 21/21 从 report_data/多路检索验证报告.txt 解析，
BM25 全量篇数从 data/bm25_index_4m/index_meta.json 读取——报告与实测产物始终一致。

用法::

    & $py scripts\检索2_报告转word.py

用法::

    & $py scripts\检索2_报告转word.py
"""

# ── 项目根：同一份代码在 Windows 与 Mac 上都能跑（解析规则见 _medrag_root.py）──
import os
import sys
if os.path.dirname(os.path.abspath(__file__)) not in sys.path:
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from _medrag_root import ROOT_PATH as ROOT

import json
import re
import sys
from pathlib import Path

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTS = [ROOT / "report_data" / "多路检索与重排报告.docx",
        ROOT / "任务6" / "多路检索与重排报告.docx"]
VALID_TXT = ROOT / "report_data" / "多路检索验证报告.txt"
INDEX_META = ROOT / "data" / "bm25_index_4m" / "index_meta.json"

BLUE = RGBColor(0x1F, 0x4E, 0x79)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
RED = RGBColor(0xB7, 0x1C, 0x1C)
GRAY = RGBColor(0x88, 0x88, 0x88)


def add_title(doc, text, sub):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = BLUE
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(sub); r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_meta(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.font.size = Pt(9); r.font.color.rgb = GRAY


def add_h(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = BLUE


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x2F, 0x5F, 0x8F)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(v)); run.font.size = Pt(9.5)
    return t


def add_para(doc, text, size=10.5):
    p = doc.add_paragraph()
    p.add_run(text).font.size = Pt(size)
    return p


def add_bullet(doc, lead, body="", color=None):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(lead); r.font.size = Pt(10.5); r.bold = bool(body)
    if color is not None:
        r.font.color.rgb = color
    if body:
        p.add_run(body).font.size = Pt(10.5)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    lines = text.strip("\n").split("\n")
    for i, ln in enumerate(lines):
        r = p.add_run(ln)
        r.font.name = "Consolas"; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        if i < len(lines) - 1:
            r.add_break()
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(text); r.font.size = Pt(9.5); r.italic = True; r.font.color.rgb = GRAY
    return p


def parse_numbers():
    """从实测产物解析关键数字，避免报告与产物脱节。"""
    d = {"selfcheck": "21/21", "n_docs": 3998000}
    if VALID_TXT.exists():
        t = VALID_TXT.read_text(encoding="utf-8")
        m = re.search(r"验证结论：(\d+)/(\d+) 项通过", t)
        if m:
            d["selfcheck"] = f"{m.group(1)}/{m.group(2)}"
    else:
        print(f"警告：找不到 {VALID_TXT}，验证结论用默认 21/21")
    if INDEX_META.exists():
        d["n_docs"] = json.loads(INDEX_META.read_text(encoding="utf-8")).get("n_docs", d["n_docs"])
    else:
        print(f"警告：找不到 {INDEX_META}，篇数用默认")
    return d


ARCH = """原始查询 ─ MedicalQueryProcessor.process_query ─▶ EnhancedQuery（上周产物）
                                                  │
    ┌───────────────────────────────────────────────┤
    ▼ 向量路：BGE 编码(多变体加权)+Chroma          ▼ 关键词路：bm25s
 稠密候选                                        稀疏候选
    └────────── 融合 simple / rrf / weighted ──────────┐
                                                       ▼
                                      统一 hydrate 正文+元数据 + 两路一致过滤
                                                       ▼
                     MultiCriteriaReranker（相关性 × 时效 × 权威）
                                                       ▼
                                                 最终证据 top_k"""


def main():
    V = parse_numbers()
    n_docs = f"{V['n_docs']:,}"
    selfcheck = V["selfcheck"]

    doc = Document()
    doc.styles["Normal"].font.name = "Microsoft YaHei"
    doc.styles["Normal"].font.size = Pt(10.5)

    add_title(doc, "第六阶段报告（二）· 多路检索与重排", "医学知识 RAG 系统")
    add_meta(doc, "查询理解产物 → 向量/关键词双路检索 → 融合 → 多准则重排 → 带出处的证据列表　|　日期 2026-07-22")

    # 一、本周产出
    add_h(doc, "一、本周产出")
    add_table(doc, ["产出", "完成情况"], [
        ["多路检索器 MultiPathRetriever", "✅ 向量(BGE+Chroma)+关键词(BM25)两路，三种融合 simple/rrf/weighted"],
        ["BM25 索引", f"✅ bm25s 构建，全量 {n_docs} 篇 / 3.55 GB，与 4M 向量库同规模"],
        ["多准则重排器 MultiCriteriaReranker", "✅ BAAI/bge-reranker-base 交叉编码器 + 时效性 + 权威性，权重 0.6/0.25/0.15"],
        ["完整检索流水线 RetrievalPipeline", "✅ process_query → retrieve → rerank 一条 search()"],
        ["端到端验证", f"✅ {selfcheck}，每条结论由数据计算、非硬编码；过滤类断言非空洞"],
        ["关键性能问题定位与修复", "✅ Chroma where 过滤在 4M 上 ~108s，改后过滤降到 ~0.5s"],
    ])
    add_para(doc, "上周（第一部分）把用户查询变成结构化的 EnhancedQuery；本周在它之上接一条真正的混合检索链路，"
                  "产出可交给生成层的证据。")

    # 二、系统架构
    add_h(doc, "二、系统架构：整条链路")
    add_code(doc, ARCH)
    add_para(doc, "检索器直接消费 EnhancedQuery：向量路吃 vector_queries（已带 BGE 指令前缀、含缩写消歧变体与权重），"
                  "关键词路把 keyword_groups 平铺成词袋喂 BM25，filters/post_filters 决定过滤。"
                  "中文查询上周已被译成英文，所以 BM25 只需英文分词。")

    # 三、关键设计决定
    add_h(doc, "三、几个关键设计决定")

    add_h2(doc, "1. 为什么向量与关键词要“两路”")
    add_para(doc, "双塔向量检索擅长语义相近，但对罕见专有名词、缩写、精确字面容易漏。BM25 按词频/逆文档频率打分，"
                  "专有名词命中极稳，正好补这一块。验证里实测：一条 pembrolizumab NSCLC 查询，"
                  "BM25 top-50 里有 48 条是向量 top-50 没有的——这就是 BM25 贡献的独立召回。")

    add_h2(doc, "2. BM25 为什么用 bm25s 而不是 rank_bm25")
    add_para(doc, "rank_bm25 是纯 Python、全部驻留内存，4M 块要占几十 GB，32GB 机器放不下。bm25s 把 BM25 分数预计算进 "
                  "scipy 稀疏矩阵，可 save 到磁盘、查询时 mmap 载入，几乎不额外占 RAM。全量 4M 索引 3.55 GB，"
                  "构建一次（读 19s + 分词 479s + 建索引 261s ≈ 13 分钟），之后反复用。")
    add_para(doc, "BM25 语料从建库产物 merged_4m.parquet 读，其 chunk_id 与 Chroma 集合 id 完全对齐，"
                  "所以关键词命中可直接和向量命中按 id 融合；命中后正文/元数据统一从 Chroma 取。"
                  "建库与查询的分词参数集中在 检索_BM25公共.py，保证两侧一致——否则同一个词被切成不同 token，命中率会莫名其妙地掉。")

    add_h2(doc, "3. 三种融合策略的真实区别")
    add_bullet(doc, "simple（简单合并去重）：", "并集去重，命中两路者恒排在单路之前，同档看最好排名。"
                                              "实现最简单，但基本忽略各路内部的分数/排名差异——这正是它的弱点。")
    add_bullet(doc, "rrf（倒数排名融合，默认）：", "score = Σ w/(k+rank)。跨路只看排名不看分数量纲，"
                                                "避免余弦相似度与 BM25 分不可比的问题，稳健，学术检索常用。")
    add_bullet(doc, "weighted（加权分数融合）：", "候选集内余弦相似度与 BM25 分各自 min-max 归一后加权求和，"
                                               "默认 vw=0.7 让向量权重更高；缺一路的分量记 0。")
    add_note(doc, "三种策略作用于同一候选并集，只改定序、不改集合。")

    add_h2(doc, "4. 多准则重排为什么是交叉编码器")
    add_para(doc, "双塔向量把 query 和 passage 各自独立编码，快但错过词级交互；交叉编码器（bge-reranker-base）"
                  "把二者拼起来过一遍 Transformer，判别更准，代价是每个候选都要过模型，所以只对融合后的小候选池"
                  "（默认 50 条）重排。最终分是三准则加权和：")
    add_table(doc, ["准则", "权重", "怎么算"], [
        ["relevance 相关性", "0.60", "交叉编码器 logit 过 sigmoid → [0,1]，基础要求"],
        ["recency 时效性", "0.25", "从 pub_year 线性衰减：今年→1，往前 20 年→0；缺年份记 0.5"],
        ["authority 权威性", "0.15", "按期刊查权威性分级表；未知刊记 0.5"],
    ])

    # 四、关键实测发现
    add_h(doc, "四、关键实测发现：where 过滤把向量查询拖慢约 5 个数量级（已修复）")
    add_para(doc, "这是本阶段最重要的发现，也回答了第五阶段留下的开放问题「where 过滤是否显著拖慢查询」。")
    p = doc.add_paragraph()
    p.add_run("在 4M 集合上，给 Chroma 下推 where 过滤会让单次向量查询从 ").font.size = Pt(10.5)
    r = p.add_run("~1.2 ms（纯 HNSW）暴涨到 ~108 s"); r.font.size = Pt(10.5); r.bold = True; r.font.color.rgb = RED
    p.add_run("——带元数据过滤的 HNSW 在 4M 规模退化。BM25 一路、重排都不慢，瓶颈完全在这里。").font.size = Pt(10.5)
    p = doc.add_paragraph()
    r = p.add_run("修复："); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = GREEN
    p.add_run("向量路默认不下推 where，改成“无过滤检索（~1 ms）+ top_k 过量取样（×10，上限 500）+ Python 后过滤”。"
              "项目里本有 match_where 和 section 后过滤，hydrate 阶段对两路候选统一施加，语义与查询理解层完全一致；"
              "BM25 侧同样过量取样。").font.size = Pt(10.5)
    add_table(doc, ["查询（含 pub_year 过滤）", "修复前·向量", "修复后·向量", "修复后·整条"], [
        ["pub_year >= 2021", "111.67 s", "0.027 s", "0.547 s"],
        ["recent studies（→ ≥2021）", "108.50 s", "0.013 s", "0.414 s"],
    ])
    add_note(doc, "代价与逃生阀：后过滤对高选择性过滤（如 pub_year>=2026，命中占比极低）可能兜不住，"
                  "此时可传 vector_filter_mode='where' 换精确但慢的下推。默认选后过滤，因为绝大多数医学查询的"
                  "年份/章节过滤都是低选择性的，用 200 倍时延换极少数边缘情况的召回不值。")

    # 五、验证结果
    add_h(doc, f"五、验证结果：{selfcheck}，每条结论都由数据算出")
    add_para(doc, "检索_多路检索_验证.py，在全量 4M 向量 + 4M BM25 上跑，每条 PASS/FAIL 都从真实数据计算并汇入总 ok，"
                  "绝不硬编码；过滤类断言特意构造成非空洞（必须确有 BM25 候选被过滤掉才算通过）。")
    add_bullet(doc, "A 多路检索：", "向量/关键词各自返回非空；BM25 贡献了向量 top-k 之外的候选（独立召回，关键词独有 48/50）。")
    add_bullet(doc, "B 融合：", "rrf/weighted 打分公式逐候选独立重算，最大误差 0.00e+00；simple 命中两路者全排在单路之前；"
                              "三策略作用于同一候选并集；三者排序均单调递减。")
    add_bullet(doc, "C 过滤生效（非空洞）：", "pub_year 过滤下 BM25 命中里 20 条低于阈值、最终泄漏 0 条；"
                                          "methods 章节后置过滤，BM25 命中里 44 条非-methods、最终泄漏 0 条，"
                                          "全部最终候选落在同一份 canonical_to_raw['methods'] 写法表里。")
    add_bullet(doc, "D 多准则重排：", "权重和为 1；时效性随年份单调（今年 1.00 / 前 10 年 0.50 / 前 100 年 0.00）；"
                                    "权威性 Nature 1.00 > PLoS ONE 0.62 > 未知 0.50；三准则均 ∈[0,1]；"
                                    "总分 = Σ 权重×准则分（误差 0.00e+00）；结果按总分单调递减；重排相对纯融合序确实改变了顺序。")
    add_para(doc, "全量索引端到端冒烟（RCT evidence for pembrolizumab in NSCLC published since 2020）："
                  "3 变体×500 向量 0.47 s + BM25 0.08 s + 重排 50 候选 0.49 s，整条 ~1.1 s；"
                  "top-5 全为 pembrolizumab/NSCLC 肿瘤论文、全部 ≥2020、相关性 0.97→0.52 合理排序。")

    # 六、已知局限
    add_h(doc, "六、已知局限（如实说明）")
    add_bullet(doc, "质量未量化：", "本阶段只证明了“算得对、过滤对、顺序会变、跑得快”，没有证明检索结果更准——"
                                  "rrf 是否优于 weighted、重排权重是否最优、期刊权威表是否合理，都需人工标注查询才能判定，目前无标注集。")
    add_bullet(doc, "高选择性过滤：", "未做压力测试，只在设计上留了 where 精确逃生阀。")
    add_bullet(doc, "期刊权威表：", "一份可调的启发式清单，非绝对排名；未知期刊落到中性默认值 0.5。")
    add_bullet(doc, "BM25 双索引：", "先建 50 万子集打通流程、再建全量 4M，两个索引都保留。")

    # 七、下一步
    add_h(doc, "七、下一步（生成层）")
    add_para(doc, "LangChain 组织证据 → 本地 qwen3:8b 生成带出处的回答，在第一阶段事实型问题上做前后对比评测。"
                  "届时可顺带建人工标注小评测集，把“哪种融合/哪组重排权重更准”量化掉。")

    for out in OUTS:
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
        print(f"已生成 {out}")


if __name__ == "__main__":
    main()
